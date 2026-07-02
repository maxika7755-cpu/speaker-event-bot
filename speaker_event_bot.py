"""
Telegram-бот для управления спикерами мероприятий.

Один файл для MVP/первого продакшена:
- aiogram 3.x: команды, inline-кнопки, FSM-мастера;
- SQLAlchemy async: SQLite по умолчанию, PostgreSQL через DATABASE_URL;
- APScheduler: напоминания о выступлениях, документах и цветах;
- OpenAI API: распознавание расписаний и обычных текстовых команд;
- загрузка расписаний: xlsx/xls/csv/txt/docx/pdf/jpg/png, с OCR при наличии pytesseract.

Переменные окружения в .env:
BOT_TOKEN=123:...
OPENAI_API_KEY=sk-...
DATABASE_URL=sqlite+aiosqlite:///speaker_bot.db
ALLOWED_USER_IDS=123456789,987654321
OPENAI_MODEL=gpt-4.1-mini
AI_PROVIDER=openai
AI_BASE_URL=
AI_API_KEY=
AI_API_STYLE=responses
DEFAULT_TIMEZONE=Europe/Moscow

Для стороннего OpenAI-compatible шлюза обычно ставьте:
AI_PROVIDER=custom
AI_API_STYLE=chat
AI_BASE_URL=https://your-gateway.example/v1
AI_API_KEY=your-gateway-key
AI_MODEL=provider-model-name
"""

from __future__ import annotations

import asyncio
import json
import logging
import os
import re
import shutil
import tempfile
from collections import defaultdict
from dataclasses import dataclass
from datetime import date, datetime, time, timedelta
from pathlib import Path
from types import SimpleNamespace
from typing import Any, Iterable
from zoneinfo import ZoneInfo

from aiogram import Bot, Dispatcher, F, Router
from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode
from aiogram.filters import CommandStart
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.dispatcher.middlewares.base import BaseMiddleware
from aiogram.types import CallbackQuery, Document, FSInputFile, Message, PhotoSize
from aiogram.utils.keyboard import InlineKeyboardBuilder
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from dotenv import load_dotenv
from openai import AsyncOpenAI
from sqlalchemy import (
    JSON,
    Boolean,
    Date,
    DateTime,
    Float,
    ForeignKey,
    Integer,
    String,
    Text,
    Time,
    func,
    select,
    text as sql_text,
)
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine
from sqlalchemy.orm import DeclarativeBase, Mapped, mapped_column, relationship, selectinload

try:
    import pandas as pd
except Exception:
    pd = None

try:
    import openpyxl
    from openpyxl.utils import get_column_letter
except Exception:
    openpyxl = None
    get_column_letter = None

try:
    import docx
except Exception:
    docx = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import pytesseract
    from PIL import Image
except Exception:
    pytesseract = None
    Image = None


load_dotenv()
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger("speaker-event-bot")

BOT_TOKEN = os.getenv("BOT_TOKEN", "").strip()
AI_PROVIDER = os.getenv("AI_PROVIDER", "openai").strip().lower()
OPENAI_API_KEY = (
    os.getenv("AI_API_KEY", "").strip()
    or os.getenv("OPENAI_API_KEY", "").strip()
    or os.getenv("FIREWORKS_API_KEY", "").strip()
)
OPENAI_MODEL = os.getenv("AI_MODEL", "").strip() or os.getenv("OPENAI_MODEL", "gpt-4.1-mini").strip()
AI_BASE_URL = (
    os.getenv("AI_BASE_URL", "").strip()
    or os.getenv("OPENAI_BASE_URL", "").strip()
    or os.getenv("FIREWORKS_BASE_URL", "").strip()
)
if not AI_BASE_URL and (AI_PROVIDER == "fireworks" or os.getenv("FIREWORKS_API_KEY")):
    AI_BASE_URL = "https://api.fireworks.ai/inference/v1"
AI_API_STYLE = os.getenv("AI_API_STYLE", "").strip().lower() or ("chat" if AI_BASE_URL else "responses")
DEFAULT_TZ = os.getenv("DEFAULT_TIMEZONE", "Europe/Moscow").strip()
ALLOWED_USER_IDS = {
    int(x.strip())
    for x in os.getenv("ALLOWED_USER_IDS", "").split(",")
    if x.strip().isdigit()
}

DATABASE_URL = os.getenv("DATABASE_URL", "sqlite+aiosqlite:///speaker_bot.db").strip()
if DATABASE_URL.startswith("postgres://"):
    DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql+asyncpg://", 1)
elif DATABASE_URL.startswith("postgresql://"):
    DATABASE_URL = DATABASE_URL.replace("postgresql://", "postgresql+asyncpg://", 1)
elif DATABASE_URL.startswith("sqlite:///"):
    DATABASE_URL = DATABASE_URL.replace("sqlite:///", "sqlite+aiosqlite:///", 1)

UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "uploads"))
UPLOAD_DIR.mkdir(exist_ok=True)

DOC_FIELDS = [
    ("contract_act_approved", "Договор / акт согласованы"),
    ("contract_act_signed", "Договор / акт подписаны"),
    ("invoice_or_receipt_received", "Счет / чек получен"),
    ("scans_sent", "Сканы в СЭД / по почте"),
    ("paid", "Оплачен"),
    ("originals_received", "ОРИГИНАЛЫ НА РУКАХ?"),
    ("originals_to_accounting", "Оригиналы в бухгалтерию"),
]
CONTRACT_TYPES = ["ИП", "ИП с НДС", "Самозанятый", "ГПХ"]
SETUP_HINT = "презентация, микрофон, петличка, кликер, барный стул, вода, цветы, другое"
PENDING_AI_ACTIONS: dict[int, dict[str, Any]] = {}
ACCESS_PERMISSIONS = [
    ("modules", "Модули"),
    ("module_manage", "Создание/удаление модулей"),
    ("speakers", "Спикеры"),
    ("schedule", "Расписание"),
    ("payments", "Оплаты и документы"),
    ("flowers", "Цветы"),
    ("reminders", "Напоминания"),
    ("settings", "Настройки"),
]
DEFAULT_GRANT_PERMISSIONS = ["modules", "speakers", "schedule", "payments", "flowers", "reminders"]


class ScheduleQualityError(RuntimeError):
    pass


class Base(DeclarativeBase):
    pass


class Module(Base):
    __tablename__ = "modules"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    name: Mapped[str] = mapped_column(String(255), index=True)
    start_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    end_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    dates_label: Mapped[str] = mapped_column(String(255), default="")
    location: Mapped[str] = mapped_column(String(255), default="")
    description: Mapped[str] = mapped_column(Text, default="")
    owner_user_id: Mapped[int] = mapped_column(Integer, index=True)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    updated_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    module_speakers: Mapped[list["ModuleSpeaker"]] = relationship(back_populates="module", cascade="all, delete-orphan")
    schedule_items: Mapped[list["ScheduleItem"]] = relationship(back_populates="module", cascade="all, delete-orphan")


class Speaker(Base):
    __tablename__ = "speakers"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    full_name: Mapped[str] = mapped_column(String(255), index=True)
    gender: Mapped[str] = mapped_column(String(32), default="")
    phone: Mapped[str] = mapped_column(String(64), default="")
    telegram: Mapped[str] = mapped_column(String(128), default="")
    email: Mapped[str] = mapped_column(String(255), default="")
    organization: Mapped[str] = mapped_column(String(255), default="")
    position: Mapped[str] = mapped_column(String(255), default="")
    comment: Mapped[str] = mapped_column(Text, default="")
    owner_user_id: Mapped[int] = mapped_column(Integer, index=True)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    updated_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    module_links: Mapped[list["ModuleSpeaker"]] = relationship(back_populates="speaker", cascade="all, delete-orphan")


class ModuleSpeaker(Base):
    __tablename__ = "module_speakers"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    module_id: Mapped[int] = mapped_column(ForeignKey("modules.id", ondelete="CASCADE"), index=True)
    speaker_id: Mapped[int] = mapped_column(ForeignKey("speakers.id", ondelete="CASCADE"), index=True)
    topic: Mapped[str] = mapped_column(String(500), default="")
    is_paid: Mapped[bool] = mapped_column(Boolean, default=False)
    amount: Mapped[float] = mapped_column(Float, default=0)
    contract_type: Mapped[str] = mapped_column(String(64), default="")
    setup: Mapped[list[str]] = mapped_column(JSON, default=list)
    flower_required: Mapped[bool] = mapped_column(Boolean, default=False)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    updated_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    module: Mapped[Module] = relationship(back_populates="module_speakers")
    speaker: Mapped[Speaker] = relationship(back_populates="module_links")
    document_status: Mapped["DocumentStatus"] = relationship(back_populates="module_speaker", cascade="all, delete-orphan")


class ScheduleItem(Base):
    __tablename__ = "schedule_items"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    module_id: Mapped[int] = mapped_column(ForeignKey("modules.id", ondelete="CASCADE"), index=True)
    speaker_id: Mapped[int | None] = mapped_column(ForeignKey("speakers.id", ondelete="SET NULL"), nullable=True, index=True)
    date: Mapped[date] = mapped_column(Date, index=True)
    start_time: Mapped[time | None] = mapped_column(Time, nullable=True)
    end_time: Mapped[time | None] = mapped_column(Time, nullable=True)
    title: Mapped[str] = mapped_column(String(500), default="")
    location: Mapped[str] = mapped_column(String(255), default="")
    format: Mapped[str] = mapped_column(String(128), default="")
    comment: Mapped[str] = mapped_column(Text, default="")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    updated_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    module: Mapped[Module] = relationship(back_populates="schedule_items")
    speaker: Mapped[Speaker | None] = relationship()


class DocumentStatus(Base):
    __tablename__ = "document_statuses"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    module_speaker_id: Mapped[int] = mapped_column(ForeignKey("module_speakers.id", ondelete="CASCADE"), unique=True)
    contract_act_approved: Mapped[bool] = mapped_column(Boolean, default=False)
    contract_act_signed: Mapped[bool] = mapped_column(Boolean, default=False)
    invoice_or_receipt_received: Mapped[bool] = mapped_column(Boolean, default=False)
    scans_sent: Mapped[bool] = mapped_column(Boolean, default=False)
    paid: Mapped[bool] = mapped_column(Boolean, default=False)
    originals_received: Mapped[bool] = mapped_column(Boolean, default=False)
    originals_to_accounting: Mapped[bool] = mapped_column(Boolean, default=False)
    progress_percent: Mapped[int] = mapped_column(Integer, default=0)
    updated_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    module_speaker: Mapped[ModuleSpeaker] = relationship(back_populates="document_status")


class Reminder(Base):
    __tablename__ = "reminders"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    module_id: Mapped[int] = mapped_column(ForeignKey("modules.id", ondelete="CASCADE"), index=True)
    speaker_id: Mapped[int | None] = mapped_column(ForeignKey("speakers.id", ondelete="SET NULL"), nullable=True, index=True)
    schedule_item_id: Mapped[int | None] = mapped_column(ForeignKey("schedule_items.id", ondelete="CASCADE"), nullable=True)
    reminder_type: Mapped[str] = mapped_column(String(64), index=True)
    remind_at: Mapped[datetime] = mapped_column(DateTime, index=True)
    is_sent: Mapped[bool] = mapped_column(Boolean, default=False)
    target_user_id: Mapped[int] = mapped_column(Integer, index=True)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    sent_at: Mapped[datetime | None] = mapped_column(DateTime, nullable=True)


class UploadedFile(Base):
    __tablename__ = "uploaded_files"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    module_id: Mapped[int] = mapped_column(ForeignKey("modules.id", ondelete="CASCADE"), index=True)
    file_name: Mapped[str] = mapped_column(String(255))
    file_type: Mapped[str] = mapped_column(String(32))
    file_path: Mapped[str] = mapped_column(String(500))
    parsed_json: Mapped[dict[str, Any]] = mapped_column(JSON, default=dict)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)


class PrettyScheduleImage(Base):
    __tablename__ = "pretty_schedule_images"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    owner_user_id: Mapped[int] = mapped_column(Integer, index=True)
    day_label: Mapped[str] = mapped_column(String(255), index=True)
    schedule_date: Mapped[date | None] = mapped_column(Date, nullable=True, index=True)
    file_name: Mapped[str] = mapped_column(String(255))
    file_path: Mapped[str] = mapped_column(String(500))
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)


class ActionLog(Base):
    __tablename__ = "action_logs"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    user_id: Mapped[int] = mapped_column(Integer, index=True)
    action_type: Mapped[str] = mapped_column(String(100))
    entity_type: Mapped[str] = mapped_column(String(100))
    entity_id: Mapped[int | None] = mapped_column(Integer, nullable=True)
    old_value: Mapped[dict[str, Any] | None] = mapped_column(JSON, nullable=True)
    new_value: Mapped[dict[str, Any] | None] = mapped_column(JSON, nullable=True)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)


class AppSetting(Base):
    __tablename__ = "settings"

    user_id: Mapped[int] = mapped_column(Integer, primary_key=True)
    timezone: Mapped[str] = mapped_column(String(64), default=DEFAULT_TZ)
    performance_reminder_minutes: Mapped[int] = mapped_column(Integer, default=40)
    docs_reminder_minutes: Mapped[int] = mapped_column(Integer, default=30)
    flowers_reminder_minutes: Mapped[int] = mapped_column(Integer, default=30)
    ai_enabled: Mapped[bool] = mapped_column(Boolean, default=True)
    post_module_reminders_enabled: Mapped[bool] = mapped_column(Boolean, default=True)
    post_module_interval_days: Mapped[int] = mapped_column(Integer, default=2)


class AccessGrant(Base):
    __tablename__ = "access_grants"

    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    owner_user_id: Mapped[int] = mapped_column(Integer, index=True)
    allowed_user_id: Mapped[int | None] = mapped_column(Integer, nullable=True, index=True)
    username: Mapped[str] = mapped_column(String(128), index=True)
    permissions: Mapped[list[str]] = mapped_column(JSON, default=list)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)


engine = create_async_engine(DATABASE_URL, echo=False)
SessionLocal = async_sessionmaker(engine, expire_on_commit=False)
router = Router()


def extract_json_text(content: str) -> str:
    content = (content or "").strip()
    fenced = re.search(r"```(?:json)?\s*(.*?)```", content, flags=re.IGNORECASE | re.DOTALL)
    if fenced:
        return fenced.group(1).strip()
    obj_start = content.find("{")
    obj_end = content.rfind("}")
    arr_start = content.find("[")
    arr_end = content.rfind("]")
    if obj_start != -1 and obj_end > obj_start and (arr_start == -1 or obj_start < arr_start):
        return content[obj_start : obj_end + 1].strip()
    if arr_start != -1 and arr_end > arr_start:
        return content[arr_start : arr_end + 1].strip()
    return content


def normalize_chat_messages(input_items: list[dict[str, Any]] | None) -> list[dict[str, str]]:
    messages = []
    for item in input_items or []:
        content = item.get("content", "")
        if isinstance(content, list):
            parts = []
            for part in content:
                if isinstance(part, dict):
                    parts.append(str(part.get("text") or part.get("content") or ""))
                else:
                    parts.append(str(part))
            content = "\n".join(part for part in parts if part.strip())
        messages.append({"role": item.get("role", "user"), "content": str(content)})
    return messages


def add_json_instruction(messages: list[dict[str, str]], schema: dict[str, Any]) -> list[dict[str, str]]:
    instruction = (
        "\n\nВерни только валидный JSON без markdown и пояснений. "
        "JSON должен соответствовать этой структуре: "
        f"{json.dumps(schema, ensure_ascii=False)[:12000]}"
    )
    updated = [dict(message) for message in messages]
    for message in updated:
        if message["role"] == "system":
            message["content"] += instruction
            return updated
    return [{"role": "system", "content": instruction.strip()}, *updated]


class ChatResponsesAdapter:
    def __init__(self, client: AsyncOpenAI) -> None:
        self.client = client

    async def create(
        self,
        model: str,
        input: list[dict[str, Any]] | None = None,
        text: dict[str, Any] | None = None,
        **kwargs: Any,
    ) -> SimpleNamespace:
        messages = normalize_chat_messages(input)
        fmt = (text or {}).get("format", {})
        name = fmt.get("name") or "json_response"
        schema = fmt.get("schema") or {"type": "object"}
        temperature = kwargs.get("temperature", 0)

        try:
            response = await self.client.chat.completions.create(
                model=model,
                messages=messages,
                response_format={
                    "type": "json_schema",
                    "json_schema": {"name": name, "schema": schema, "strict": True},
                },
                temperature=temperature,
            )
        except Exception as exc:
            error_text = str(exc).lower()
            if "json_schema" not in error_text and "response_format" not in error_text and "schema" not in error_text:
                raise
            try:
                response = await self.client.chat.completions.create(
                    model=model,
                    messages=add_json_instruction(messages, schema),
                    response_format={"type": "json_object"},
                    temperature=temperature,
                )
            except Exception as second_exc:
                second_error_text = str(second_exc).lower()
                if "response_format" not in second_error_text and "json_object" not in second_error_text:
                    raise
                response = await self.client.chat.completions.create(
                    model=model,
                    messages=add_json_instruction(messages, schema),
                    temperature=temperature,
                )

        content = response.choices[0].message.content or "{}"
        return SimpleNamespace(output_text=extract_json_text(content))


class AIClientAdapter:
    def __init__(self, client: AsyncOpenAI) -> None:
        self.chat = client.chat
        self.responses = ChatResponsesAdapter(client) if AI_API_STYLE == "chat" else client.responses


_openai_sdk_client = AsyncOpenAI(api_key=OPENAI_API_KEY, base_url=AI_BASE_URL or None) if OPENAI_API_KEY else None
openai_client = AIClientAdapter(_openai_sdk_client) if _openai_sdk_client else None


class AccessMiddleware(BaseMiddleware):
    async def __call__(self, handler: Any, event: Any, data: dict[str, Any]) -> Any:
        user = data.get("event_from_user")
        if not user:
            return await handler(event, data)
        if await access_allowed(user.id, user.username):
            if isinstance(event, CallbackQuery) and not await callback_allowed(user.id, user.username, event.data or ""):
                await event.answer("⛔️ Эта кнопка для вас отключена.", show_alert=True)
                return None
            return await handler(event, data)
        if isinstance(event, Message):
            await event.answer("⛔️ Доступ к боту закрыт.")
        elif isinstance(event, CallbackQuery):
            await event.answer("⛔️ Доступ закрыт.", show_alert=True)
        return None


async def ai_json_create(system_prompt: str, user_prompt: str, schema: dict[str, Any], name: str) -> dict[str, Any]:
    if not openai_client or not OPENAI_API_KEY:
        raise RuntimeError("AI API key is not configured")

    response = await openai_client.responses.create(
        model=OPENAI_MODEL,
        input=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        text={"format": {"type": "json_schema", "name": name, "schema": schema, "strict": True}},
    )
    return json.loads(response.output_text)


async def ai_text_create(system_prompt: str, user_prompt: str) -> str:
    if not _openai_sdk_client or not OPENAI_API_KEY:
        raise RuntimeError("AI API key is not configured")
    if AI_API_STYLE == "chat":
        response = await _openai_sdk_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.1,
        )
        return (response.choices[0].message.content or "").strip()
    response = await _openai_sdk_client.responses.create(
        model=OPENAI_MODEL,
        input=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )
    return response.output_text.strip()


class ModuleCreate(StatesGroup):
    name = State()
    dates = State()
    location = State()
    description = State()


class ModuleEdit(StatesGroup):
    field = State()


class ModuleAIQuestion(StatesGroup):
    question = State()


class SpeakerCreate(StatesGroup):
    full_name = State()
    gender = State()
    phone = State()
    telegram = State()
    email = State()
    organization = State()
    position = State()
    topic = State()
    is_paid = State()
    amount = State()
    contract_type = State()
    setup = State()
    comment = State()


class SpeakerEdit(StatesGroup):
    setup = State()
    amount = State()
    data = State()


class ScheduleUpload(StatesGroup):
    waiting_file = State()


class PrettyScheduleUpload(StatesGroup):
    day_label = State()
    image = State()


class ManualSchedule(StatesGroup):
    item_date = State()
    start_time = State()
    end_time = State()
    title = State()
    speaker_name = State()
    location = State()
    item_format = State()
    comment = State()


class SettingsEdit(StatesGroup):
    timezone = State()
    performance_minutes = State()
    docs_minutes = State()
    flowers_minutes = State()
    post_interval = State()


class AccessEdit(StatesGroup):
    username = State()


def html_escape(text: Any) -> str:
    s = "" if text is None else str(text)
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def money(value: float | int | None) -> str:
    if not value:
        return "0 руб."
    return f"{int(value):,}".replace(",", " ") + " руб."


def parse_time(value: str) -> time | None:
    value = (value or "").strip()
    m = re.search(r"(\d{1,2})[:.](\d{2})", value)
    if not m:
        return None
    h, minute = int(m.group(1)), int(m.group(2))
    if 0 <= h <= 23 and 0 <= minute <= 59:
        return time(h, minute)
    return None


def parse_date_value(value: str, fallback_year: int | None = None) -> date | None:
    value = (value or "").strip().lower()
    fallback_year = fallback_year or datetime.now().year
    months = {
        "январ": 1, "феврал": 2, "март": 3, "апрел": 4, "ма": 5, "июн": 6,
        "июл": 7, "август": 8, "сентябр": 9, "октябр": 10, "ноябр": 11, "декабр": 12,
    }
    for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%d.%m.%y", "%d.%m"):
        try:
            parsed = datetime.strptime(value, fmt)
            return date(fallback_year if fmt == "%d.%m" else parsed.year, parsed.month, parsed.day)
        except ValueError:
            pass
    m = re.search(r"(\d{1,2})\s+([а-яё]+)(?:\s+(\d{4}))?", value)
    if m:
        month = next((num for key, num in months.items() if m.group(2).startswith(key)), None)
        if month:
            return date(int(m.group(3) or fallback_year), month, int(m.group(1)))
    return None


def parse_date_range(text: str) -> tuple[date | None, date | None, str]:
    label = text.strip()
    year = datetime.now().year
    dates = re.findall(r"\d{4}-\d{2}-\d{2}|\d{1,2}\.\d{1,2}(?:\.\d{2,4})?|\d{1,2}\s+[а-яё]+(?:\s+\d{4})?", label.lower())
    if len(dates) >= 2:
        return parse_date_value(dates[0], year), parse_date_value(dates[1], year), label
    first = parse_date_value(label, year)
    return first, first, label


def split_setup(text: str | Iterable[str] | None) -> list[str]:
    if not text:
        return []
    if isinstance(text, list):
        raw = text
    else:
        raw = re.split(r"[,;\n]+", str(text))
    return [x.strip().lower() for x in raw if x and x.strip()]


FEMALE_FIRST_NAMES = {
    "анна", "мария", "елена", "евгения", "светлана", "татьяна", "ольга", "наталья",
    "екатерина", "александра", "юлия", "дарья", "ирина", "ксения", "валерия",
    "анастасия", "полина", "алина", "виктория", "любовь", "надежда",
    "елизавета", "марина", "галина", "лариса", "вероника", "яна", "нина",
    "диана", "карина", "жанна", "оксана", "алёна", "алена", "милана",
}

NON_SPEAKER_WORDS = {
    "завтрак", "обед", "ужин", "кофе-брейк", "перерыв", "дорога", "рефлексия",
    "открытие", "знакомство", "жога", "лпр", "вк", "сенеж", "сообщество",
    "сбор гостей", "день в рбк", "дорога в рбк", "дорога в отель", "отель",
    "мастер-класс", "командообразование", "стрельба", "урок связи",
}

NON_SPEAKER_MARKERS = {
    "завтрак", "обед", "ужин", "кофе", "дорога", "сбор", "гостей", "отель",
    "день в", "рбк", "жога", "лпр", "командообраз", "стрельба", "квест",
    "станция", "станции", "регистрация", "перерыв", "заселение", "выезд",
}

ACTIVITY_CONTINUATION_MARKERS = {
    "станц", "стрельб", "урок", "связ", "квест", "лазер", "тагар", "полоса",
    "препятств", "оруж", "городк", "команд", "упражнен", "практик",
}


def infer_gender_from_name(full_name: str) -> str:
    parts = [p.strip(" ,.") for p in full_name.lower().split() if p.strip(" ,.")]
    if not parts:
        return ""
    if any(p in FEMALE_FIRST_NAMES for p in parts):
        return "женщина"
    if any(p.endswith(("вна", "чна", "инична")) for p in parts):
        return "женщина"
    first = parts[0]
    if first.endswith(("ова", "ева", "ина", "ая", "ская", "цкая")):
        return "женщина"
    if len(parts) > 1 and parts[1].endswith(("а", "я")):
        return "женщина"
    return ""


def split_speaker_names(value: Any) -> list[str]:
    if not value:
        return []
    if isinstance(value, list):
        raw = []
        for item in value:
            raw.extend(split_speaker_names(item))
        return unique_names(raw)
    text = re.sub(r"\s+", " ", str(value)).strip()
    if not text:
        return []
    text = re.sub(r"(?i)\bспикер(?:ы)?\s*:\s*", "", text)
    parts = re.split(r"\s*(?:;|,|\n| и |\+|/)\s*", text)
    return unique_names(clean_speaker_name(p) for p in parts)


def clean_speaker_name(name: str) -> str:
    name = re.sub(r"\([^)]*\)", "", name or "")
    name = re.sub(r"\b(эксперт|модератор|автор|создатель|основатель|президент|победитель|предприниматель)\b.*", "", name, flags=re.I)
    name = name.strip(" -–—,.;:|\t")
    words = [w for w in name.split() if w]
    if len(words) > 4:
        words = words[:3]
    candidate = " ".join(words)
    low = candidate.lower()
    if not candidate or low in NON_SPEAKER_WORDS:
        return ""
    if len(candidate) < 5:
        return ""
    return candidate


def unique_names(names: Iterable[str]) -> list[str]:
    result = []
    seen = set()
    for name in names:
        cleaned = clean_speaker_name(str(name))
        key = cleaned.lower()
        if cleaned and key not in seen:
            result.append(cleaned)
            seen.add(key)
    return result


def looks_like_speaker_name(text: str) -> bool:
    text = clean_speaker_name(text)
    if not text:
        return False
    low = text.lower()
    if low in NON_SPEAKER_WORDS or any(marker in low for marker in NON_SPEAKER_MARKERS):
        return False
    words = text.split()
    capitalized = sum(1 for w in words if w[:1].isupper())
    has_known_first_name = any(word.lower().strip(".,") in FEMALE_FIRST_NAMES for word in words)
    return (2 <= len(words) <= 4 and capitalized >= 2) or has_known_first_name or infer_gender_from_name(text) == "женщина"


def extract_speaker_candidates(text: str) -> list[str]:
    if not text:
        return []
    candidates = []
    role_pattern = re.compile(
        r"(?:спикер|эксперт|ведущий|модератор|тренер|лектор|фасилитатор)\s*:?\s*"
        r"([А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\s+[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?(?:\s+[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?)?)",
        flags=re.I,
    )
    for match in role_pattern.finditer(text):
        candidate = clean_speaker_name(match.group(1))
        if looks_like_speaker_name(candidate):
            candidates.append(candidate)
    pattern = re.compile(
        r"\b[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\s+"
        r"[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?"
        r"(?:\s+[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?)?"
    )
    for match in pattern.finditer(text):
        candidate = clean_speaker_name(match.group(0))
        if looks_like_speaker_name(candidate):
            candidates.append(candidate)
    return unique_names(candidates)


def remove_speaker_names_from_title(title: str, speakers: Iterable[str]) -> str:
    cleaned = title or ""
    for speaker in speakers:
        cleaned = re.sub(re.escape(speaker), "", cleaned).strip(" -–—,.;:/")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned or title


def is_activity_continuation(current_title: str, next_text: str) -> bool:
    if not current_title or not next_text:
        return False
    low_current = current_title.lower()
    low_next = next_text.lower()
    if low_next.startswith(("(", "-", "и ", "или ")):
        return True
    if "командообраз" in low_current and any(marker in low_next for marker in ACTIVITY_CONTINUATION_MARKERS):
        return True
    if "мастер-класс" in low_current and low_next[:1].islower():
        return True
    return False


def item_speaker_names(item: dict[str, Any]) -> list[str]:
    names = split_speaker_names(item.get("speaker_names") or [])
    names.extend(split_speaker_names(item.get("speaker_name") or ""))
    if not names:
        names.extend(extract_speaker_candidates(item.get("comment", "")))
    return unique_names(names)


def doc_progress(ds: DocumentStatus | None) -> int:
    if not ds:
        return 0
    done = sum(1 for field, _ in DOC_FIELDS if getattr(ds, field))
    ds.progress_percent = round(done / len(DOC_FIELDS) * 100)
    return ds.progress_percent


async def get_settings(session: AsyncSession, user_id: int) -> AppSetting:
    settings = await session.get(AppSetting, user_id)
    if not settings:
        settings = AppSetting(user_id=user_id)
        session.add(settings)
        await session.flush()
    return settings


def normalize_username(username: str | None) -> str:
    return (username or "").strip().lstrip("@").lower()


async def visible_owner_ids(session: AsyncSession, user_id: int, username: str | None = None) -> list[int]:
    owners = {user_id}
    uname = normalize_username(username)
    stmt = select(AccessGrant).where(
        (AccessGrant.allowed_user_id == user_id)
        | (AccessGrant.username == uname if uname else AccessGrant.id == -1)
    )
    grants = (await session.scalars(stmt)).all()
    for grant in grants:
        if grant.allowed_user_id is None and uname and grant.username == uname:
            grant.allowed_user_id = user_id
        owners.add(grant.owner_user_id)
    return list(owners)


def normalize_access_permissions(value: Any) -> list[str]:
    known = {key for key, _title in ACCESS_PERMISSIONS}
    if isinstance(value, str):
        try:
            value = json.loads(value)
        except json.JSONDecodeError:
            return DEFAULT_GRANT_PERMISSIONS.copy()
    if isinstance(value, list):
        return [str(item) for item in value if str(item) in known]
    return DEFAULT_GRANT_PERMISSIONS.copy()


async def access_permissions(session: AsyncSession, user_id: int, username: str | None = None) -> set[str] | None:
    if not ALLOWED_USER_IDS or user_id in ALLOWED_USER_IDS:
        return None
    uname = normalize_username(username)
    stmt = select(AccessGrant).where(
        (AccessGrant.allowed_user_id == user_id)
        | (AccessGrant.username == uname if uname else AccessGrant.id == -1)
    )
    grants = (await session.scalars(stmt)).all()
    if not grants:
        return set()
    result: set[str] = set()
    for grant in grants:
        if grant.allowed_user_id is None and uname and grant.username == uname:
            grant.allowed_user_id = user_id
        result.update(normalize_access_permissions(grant.permissions))
    return result


async def effective_permissions(user_id: int, username: str | None = None) -> set[str] | None:
    async with SessionLocal() as session:
        permissions = await access_permissions(session, user_id, username)
        await session.commit()
        return permissions


def callback_permission_key(callback_data: str) -> str | None:
    data = callback_data or ""
    if data in {"menu"}:
        return None
    if data in {"modules"} or data.startswith("module:"):
        return "modules"
    if data.startswith("module_ai_question:") or data == "ai_module_menu":
        return "modules"
    if data in {"module_create"} or data.startswith(("module_edit:", "module_delete_ask:", "module_delete:")):
        return "module_manage"
    if data in {"speakers"} or data.startswith(("speaker:", "speaker_create:", "speaker_edit_", "speaker_field:", "speaker_delete", "module_speakers:", "module_speakers_delete_all", "speakers_delete_all")):
        return "speakers"
    if data == "schedule_all" or data.startswith(("schedule_", "upload_schedule:", "pretty_schedule")):
        return "schedule"
    if data == "payments_all" or data.startswith(("module_payments:", "docs:", "doc_toggle:")):
        return "payments"
    if data == "flowers_all" or data.startswith("module_flowers:"):
        return "flowers"
    if data == "reminders" or data.startswith("recreate_reminders:"):
        return "reminders"
    if data.startswith("access_") or data.startswith("settings") or data == "settings":
        return "settings"
    return None


async def callback_allowed(user_id: int, username: str | None, callback_data: str) -> bool:
    permissions = await effective_permissions(user_id, username)
    if permissions is None:
        return True
    key = callback_permission_key(callback_data)
    return key is None or key in permissions


async def log_action(
    session: AsyncSession,
    user_id: int,
    action_type: str,
    entity_type: str,
    entity_id: int | None,
    old: dict[str, Any] | None = None,
    new: dict[str, Any] | None = None,
) -> None:
    session.add(ActionLog(user_id=user_id, action_type=action_type, entity_type=entity_type, entity_id=entity_id, old_value=old, new_value=new))


def permission_allows(permissions: set[str] | None, callback_data: str) -> bool:
    if permissions is None:
        return True
    key = callback_permission_key(callback_data)
    return key is None or key in permissions


def kb_main(permissions: set[str] | None = None) -> InlineKeyboardBuilder:
    kb = InlineKeyboardBuilder()
    items = [
        ("📁 Мои модули", "modules"),
        ("➕ Создать модуль", "module_create"),
        ("👤 Спикеры", "speakers"),
        ("📅 Расписание", "schedule_all"),
        ("🖼 Красивое расписание", "pretty_schedule"),
        ("🤖 Спросить ИИ", "ai_module_menu"),
        ("💰 Оплаты и документы", "payments_all"),
        ("🌸 Цветы", "flowers_all"),
        ("🔔 Напоминания", "reminders"),
        ("⚙️ Настройки", "settings"),
    ]
    for text, data in items:
        if permission_allows(permissions, data):
            kb.button(text=text, callback_data=data)
    kb.adjust(2)
    return kb


def kb_back(module_id: int | None = None) -> InlineKeyboardBuilder:
    kb = InlineKeyboardBuilder()
    if module_id:
        kb.button(text="⬅️ Назад", callback_data=f"module:{module_id}")
    kb.button(text="🏠 Главное меню", callback_data="menu")
    kb.adjust(1)
    return kb


def kb_module(module_id: int, permissions: set[str] | None = None) -> InlineKeyboardBuilder:
    kb = InlineKeyboardBuilder()
    buttons = [
        ("👤 Спикеры модуля", f"module_speakers:{module_id}"),
        ("📅 Расписание по дням", f"schedule_days:{module_id}"),
        ("💰 Оплаты и документы", f"module_payments:{module_id}"),
        ("🌸 Цветы", f"module_flowers:{module_id}"),
        ("➕ Добавить спикера", f"speaker_create:{module_id}"),
        ("📎 Загрузить сетку", f"upload_schedule:{module_id}"),
        ("✏️ Изменить модуль", f"module_edit:{module_id}"),
        ("🗑 Удалить модуль", f"module_delete_ask:{module_id}"),
        ("🏠 Главное меню", "menu"),
    ]
    for text, data in buttons:
        if permission_allows(permissions, data):
            kb.button(text=text, callback_data=data)
    kb.adjust(2, 2, 2, 1, 2, 1, 1)
    return kb


def kb_speaker(ms_id: int, module_id: int) -> InlineKeyboardBuilder:
    kb = InlineKeyboardBuilder()
    buttons = [
        ("✏️ Изменить данные", f"speaker_edit_data:{ms_id}"),
        ("⚙️ Изменить сетап", f"speaker_edit_setup:{ms_id}"),
        ("💰 Изменить сумму", f"speaker_edit_amount:{ms_id}"),
        ("📄 Статусы документов", f"docs:{ms_id}"),
        ("📅 Перейти к расписанию", f"schedule_days:{module_id}"),
        ("🔔 Настроить напоминания", f"recreate_reminders:{ms_id}"),
        ("🗑 Удалить спикера", f"speaker_delete_ask:{ms_id}"),
        ("⬅️ К модулю", f"module:{module_id}"),
        ("🏠 Главное меню", "menu"),
    ]
    for text, data in buttons:
        kb.button(text=text, callback_data=data)
    kb.adjust(2, 2, 1, 1, 1, 2)
    return kb


async def find_speaker(session: AsyncSession, user_id: int, name: str) -> Speaker | None:
    name_norm = name.strip().lower()
    if not name_norm:
        return None
    speakers = (await session.scalars(select(Speaker).where(Speaker.owner_user_id == user_id))).all()
    for sp in speakers:
        if name_norm in sp.full_name.lower() or sp.full_name.lower() in name_norm:
            return sp
    first_token = name_norm.split()[0]
    for sp in speakers:
        if first_token and first_token in sp.full_name.lower():
            return sp
    return None


async def find_module(session: AsyncSession, user_id: int, name: str) -> Module | None:
    name_norm = name.strip().lower()
    modules = (await session.scalars(select(Module).where(Module.owner_user_id == user_id))).all()
    for module in modules:
        if name_norm in module.name.lower() or module.name.lower() in name_norm:
            return module
    return modules[0] if modules and not name_norm else None


async def get_module_speaker(session: AsyncSession, module_id: int, speaker_id: int) -> ModuleSpeaker | None:
    return await session.scalar(
        select(ModuleSpeaker)
        .options(selectinload(ModuleSpeaker.document_status), selectinload(ModuleSpeaker.speaker), selectinload(ModuleSpeaker.module))
        .where(ModuleSpeaker.module_id == module_id, ModuleSpeaker.speaker_id == speaker_id)
    )


async def ensure_module_speaker(
    session: AsyncSession,
    module: Module,
    speaker: Speaker,
    topic: str = "",
    setup: list[str] | None = None,
    is_paid: bool = False,
    amount: float = 0,
    contract_type: str = "",
) -> ModuleSpeaker:
    link = await get_module_speaker(session, module.id, speaker.id)
    if not link:
        link = ModuleSpeaker(
            module_id=module.id,
            speaker_id=speaker.id,
            topic=topic or "",
            setup=setup or [],
            is_paid=is_paid,
            amount=amount or 0,
            contract_type=contract_type or "",
            flower_required=speaker.gender.lower().startswith("жен"),
        )
        session.add(link)
        await session.flush()
    else:
        if topic:
            link.topic = topic
        if setup:
            link.setup = setup
        if is_paid:
            link.is_paid = True
            link.amount = amount or link.amount
            link.contract_type = contract_type or link.contract_type
    if link.is_paid and not link.document_status:
        session.add(DocumentStatus(module_speaker_id=link.id))
    return link


async def module_card(
    session: AsyncSession,
    module_id: int,
    permissions: set[str] | None = None,
) -> tuple[str, InlineKeyboardBuilder]:
    module = await session.get(Module, module_id)
    if not module:
        return "Модуль не найден.", kb_main()
    links = (
        await session.scalars(
            select(ModuleSpeaker)
            .where(ModuleSpeaker.module_id == module_id)
            .options(selectinload(ModuleSpeaker.speaker), selectinload(ModuleSpeaker.document_status))
        )
    ).all()
    paid = [x for x in links if x.is_paid]
    flowers = [x for x in links if x.flower_required or x.speaker.gender.lower().startswith("жен")]
    progresses = [doc_progress(x.document_status) for x in paid if x.document_status]
    avg_progress = round(sum(progresses) / len(progresses)) if progresses else 0
    now = datetime.now()
    upcoming = await session.scalar(
        select(ScheduleItem)
        .where(ScheduleItem.module_id == module_id, ScheduleItem.date >= now.date())
        .order_by(ScheduleItem.date, ScheduleItem.start_time)
        .options(selectinload(ScheduleItem.speaker))
    )
    lines = [
        f"📁 <b>Модуль:</b> {html_escape(module.name)}",
        f"📅 <b>Даты:</b> {html_escape(module.dates_label or date_range_label(module))}",
        f"📍 <b>Локация:</b> {html_escape(module.location or 'не указана')}",
        "",
        f"👤 <b>Спикеров всего:</b> {len(links)}",
        f"💰 <b>Платных спикеров:</b> {len(paid)}",
        f"🌸 <b>Цветы нужны:</b> {len(flowers)} букет(а)",
        f"📄 <b>Документы закрыты:</b> {avg_progress}%",
    ]
    if upcoming:
        speaker_name = upcoming.speaker.full_name if upcoming.speaker else "Спикер не указан"
        ms = await get_module_speaker(session, module.id, upcoming.speaker_id) if upcoming.speaker_id else None
        lines += [
            "",
            "<b>Ближайшее выступление:</b>",
            f"{time_label(upcoming.start_time)} - {html_escape(speaker_name)}",
            f"Тема: {html_escape(upcoming.title or (ms.topic if ms else ''))}",
            f"Сетап: {html_escape(', '.join(ms.setup) if ms and ms.setup else 'не указан')}",
        ]
    return "\n".join(lines), kb_module(module_id, permissions)


def date_range_label(module: Module) -> str:
    if module.start_date and module.end_date and module.start_date != module.end_date:
        return f"{module.start_date:%d.%m.%Y} - {module.end_date:%d.%m.%Y}"
    if module.start_date:
        return f"{module.start_date:%d.%m.%Y}"
    return "не указаны"


def time_label(value: time | None) -> str:
    return value.strftime("%H:%M") if value else "--:--"


def time_to_minutes(value: str | time | None) -> int:
    if isinstance(value, time):
        return value.hour * 60 + value.minute
    parsed = parse_time(str(value or ""))
    return parsed.hour * 60 + parsed.minute if parsed else 0


async def speaker_card(session: AsyncSession, ms_id: int) -> tuple[str, InlineKeyboardBuilder]:
    ms = await session.scalar(
        select(ModuleSpeaker)
        .where(ModuleSpeaker.id == ms_id)
        .options(selectinload(ModuleSpeaker.speaker), selectinload(ModuleSpeaker.module), selectinload(ModuleSpeaker.document_status))
    )
    if not ms:
        return "Спикер не найден.", kb_main()
    item = await session.scalar(
        select(ScheduleItem)
        .where(ScheduleItem.module_id == ms.module_id, ScheduleItem.speaker_id == ms.speaker_id)
        .order_by(ScheduleItem.date, ScheduleItem.start_time)
    )
    progress = doc_progress(ms.document_status)
    lines = [
        f"👤 <b>Спикер:</b> {html_escape(ms.speaker.full_name)}",
        f"📁 <b>Модуль:</b> {html_escape(ms.module.name)}",
        f"📅 <b>Дата выступления:</b> {item.date.strftime('%d.%m.%Y') if item else 'не указана'}",
        f"⏰ <b>Время:</b> {time_label(item.start_time) if item else '--:--'}-{time_label(item.end_time) if item else '--:--'}",
        f"🎤 <b>Тема:</b> {html_escape(ms.topic or (item.title if item else 'не указана'))}",
        "",
        f"📞 <b>Телефон:</b> {html_escape(ms.speaker.phone or 'не указан')}",
        f"💬 <b>Telegram:</b> {html_escape(ms.speaker.telegram or 'не указан')}",
        f"📧 <b>Email:</b> {html_escape(ms.speaker.email or 'не указан')}",
        "",
        f"💰 <b>Статус:</b> {'платный' if ms.is_paid else 'бесплатный'}",
        f"💵 <b>Сумма:</b> {money(ms.amount)}",
        f"📄 <b>Тип оформления:</b> {html_escape(ms.contract_type or 'не указан')}",
        "",
        "⚙️ <b>Сетап:</b>",
    ]
    lines += [f"- {html_escape(x)}" for x in (ms.setup or ["не указан"])]
    lines += ["", f"📄 <b>Документы:</b> {progress}%"]
    return "\n".join(lines), kb_speaker(ms.id, ms.module_id)


def split_long_text(text: str, limit: int = 3500) -> list[str]:
    text = text or ""
    chunks = []
    while len(text) > limit:
        cut = text.rfind("\n", 0, limit)
        if cut < limit // 2:
            cut = limit
        chunks.append(text[:cut].strip())
        text = text[cut:].strip()
    if text:
        chunks.append(text)
    return chunks or [""]


async def build_module_ai_context(session: AsyncSession, module_id: int) -> str:
    module = await session.get(Module, module_id)
    if not module:
        return "Модуль не найден."

    links = (
        await session.scalars(
            select(ModuleSpeaker)
            .where(ModuleSpeaker.module_id == module_id)
            .options(selectinload(ModuleSpeaker.speaker), selectinload(ModuleSpeaker.document_status))
            .order_by(ModuleSpeaker.id)
        )
    ).all()
    items = (
        await session.scalars(
            select(ScheduleItem)
            .where(ScheduleItem.module_id == module_id)
            .options(selectinload(ScheduleItem.speaker))
            .order_by(ScheduleItem.date, ScheduleItem.start_time, ScheduleItem.end_time)
        )
    ).all()

    lines = [
        f"Модуль: {module.name}",
        f"Даты: {module.dates_label or date_range_label(module)}",
        f"Локация: {module.location or 'не указана'}",
        f"Описание: {module.description or 'не указано'}",
        "",
        "СПИКЕРЫ:",
    ]
    if links:
        for ms in links:
            docs = doc_progress(ms.document_status)
            flower = "да" if ms.flower_required or ms.speaker.gender.lower().startswith("жен") else "нет"
            lines.extend(
                [
                    f"- {ms.speaker.full_name}",
                    f"  Пол: {ms.speaker.gender or 'не указан'}",
                    f"  Тема: {ms.topic or 'не указана'}",
                    f"  Телефон: {ms.speaker.phone or 'не указан'}",
                    f"  Telegram: {ms.speaker.telegram or 'не указан'}",
                    f"  Email: {ms.speaker.email or 'не указан'}",
                    f"  Организация/должность: {ms.speaker.organization or '-'} / {ms.speaker.position or '-'}",
                    f"  Оплата: {money(ms.amount) if ms.is_paid else 'не отмечена'}",
                    f"  Документы: {docs}%",
                    f"  Сетап: {', '.join(ms.setup or []) if ms.setup else 'не указан'}",
                    f"  Цветы: {flower}",
                ]
            )
    else:
        lines.append("- нет спикеров")

    lines.extend(["", "РАСПИСАНИЕ:"])
    if items:
        for item in items:
            speaker = item.speaker.full_name if item.speaker else "спикера нет"
            lines.append(
                f"- {item.date:%Y-%m-%d} {time_label(item.start_time)}-{time_label(item.end_time)} | "
                f"Спикер: {speaker} | Тема: {item.title or '-'} | "
                f"Локация: {item.location or module.location or '-'} | Формат: {item.format or '-'} | "
                f"Комментарий: {item.comment or '-'}"
            )
    else:
        lines.append("- расписание пустое")
    return "\n".join(lines)


async def send_or_edit(event: Message | CallbackQuery, text: str, kb: InlineKeyboardBuilder | None = None) -> None:
    markup = kb.as_markup() if kb else None
    if isinstance(event, CallbackQuery):
        await event.message.edit_text(text, reply_markup=markup)
        await event.answer()
    else:
        await event.answer(text, reply_markup=markup)


async def access_allowed(user_id: int, username: str | None = None) -> bool:
    if not ALLOWED_USER_IDS or user_id in ALLOWED_USER_IDS:
        return True
    async with SessionLocal() as session:
        owners = await visible_owner_ids(session, user_id, username)
        await session.commit()
    return len(set(owners) - {user_id}) > 0


@router.message(CommandStart())
async def start(message: Message, state: FSMContext) -> None:
    await state.clear()
    if not await access_allowed(message.from_user.id, message.from_user.username):
        await message.answer("⛔️ Доступ к боту закрыт. Добавьте ваш Telegram ID в ALLOWED_USER_IDS.")
        return
    async with SessionLocal() as session:
        await get_settings(session, message.from_user.id)
        permissions = await access_permissions(session, message.from_user.id, message.from_user.username)
        await session.commit()
    await message.answer(
        "Готов помогать со спикерами, расписанием, документами, оплатами и напоминаниями.",
        reply_markup=kb_main(permissions).as_markup(),
    )


@router.callback_query(F.data == "menu")
async def menu(callback: CallbackQuery, state: FSMContext) -> None:
    await state.clear()
    permissions = await effective_permissions(callback.from_user.id, callback.from_user.username)
    await send_or_edit(callback, "Главное меню", kb_main(permissions))


@router.callback_query(F.data == "module_create")
async def module_create(callback: CallbackQuery, state: FSMContext) -> None:
    await state.set_state(ModuleCreate.name)
    await callback.message.edit_text("➕ Введите название модуля:", reply_markup=kb_back().as_markup())
    await callback.answer()


@router.message(ModuleCreate.name)
async def module_create_name(message: Message, state: FSMContext) -> None:
    await state.update_data(name=message.text.strip())
    await state.set_state(ModuleCreate.dates)
    await message.answer("📅 Введите даты модуля. Например: 24-25 июня или 2026-06-24 - 2026-06-25")


@router.message(ModuleCreate.dates)
async def module_create_dates(message: Message, state: FSMContext) -> None:
    start, end, label = parse_date_range(message.text)
    await state.update_data(start_date=start.isoformat() if start else None, end_date=end.isoformat() if end else None, dates_label=label)
    await state.set_state(ModuleCreate.location)
    await message.answer("📍 Введите локацию:")


@router.message(ModuleCreate.location)
async def module_create_location(message: Message, state: FSMContext) -> None:
    await state.update_data(location=message.text.strip())
    await state.set_state(ModuleCreate.description)
    await message.answer("📝 Добавьте короткое описание или отправьте «-»:")


@router.message(ModuleCreate.description)
async def module_create_finish(message: Message, state: FSMContext) -> None:
    data = await state.get_data()
    async with SessionLocal() as session:
        module = Module(
            name=data["name"],
            start_date=date.fromisoformat(data["start_date"]) if data.get("start_date") else None,
            end_date=date.fromisoformat(data["end_date"]) if data.get("end_date") else None,
            dates_label=data.get("dates_label", ""),
            location=data.get("location", ""),
            description="" if message.text.strip() == "-" else message.text.strip(),
            owner_user_id=message.from_user.id,
        )
        session.add(module)
        await session.flush()
        await log_action(session, message.from_user.id, "create", "module", module.id, new={"name": module.name})
        await session.commit()
        text, kb = await module_card(session, module.id)
    await state.clear()
    await message.answer(text, reply_markup=kb.as_markup())


@router.callback_query(F.data == "modules")
async def modules_list(callback: CallbackQuery) -> None:
    async with SessionLocal() as session:
        owner_ids = await visible_owner_ids(session, callback.from_user.id, callback.from_user.username)
        permissions = await access_permissions(session, callback.from_user.id, callback.from_user.username)
        modules = (
            await session.scalars(
                select(Module).where(Module.owner_user_id.in_(owner_ids)).order_by(Module.start_date, Module.created_at)
            )
        ).all()
        await session.commit()
    kb = InlineKeyboardBuilder()
    for module in modules:
        kb.button(text=f"{module.name} | {module.dates_label or date_range_label(module)}", callback_data=f"module:{module.id}")
    if permission_allows(permissions, "module_create"):
        kb.button(text="➕ Создать модуль", callback_data="module_create")
    kb.button(text="🏠 Главное меню", callback_data="menu")
    kb.adjust(1)
    await send_or_edit(callback, "📁 <b>Мои модули</b>" if modules else "Пока нет модулей.", kb)


@router.callback_query(F.data.startswith("module:"))
async def module_open(callback: CallbackQuery) -> None:
    module_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        permissions = await access_permissions(session, callback.from_user.id, callback.from_user.username)
        text, kb = await module_card(session, module_id, permissions)
    await send_or_edit(callback, text, kb)


@router.callback_query(F.data == "ai_module_menu")
async def ai_module_menu(callback: CallbackQuery) -> None:
    if not openai_client or not OPENAI_API_KEY:
        await callback.answer("ИИ не подключен. Проверьте AI_API_KEY, AI_BASE_URL и AI_MODEL.", show_alert=True)
        return
    async with SessionLocal() as session:
        owner_ids = await visible_owner_ids(session, callback.from_user.id, callback.from_user.username)
        modules = (
            await session.scalars(
                select(Module).where(Module.owner_user_id.in_(owner_ids)).order_by(Module.start_date, Module.created_at)
            )
        ).all()
    kb = InlineKeyboardBuilder()
    for module in modules:
        kb.button(text=f"{module.name} | {module.dates_label or date_range_label(module)}", callback_data=f"module_ai_question:{module.id}")
    kb.button(text="🏠 Главное меню", callback_data="menu")
    kb.adjust(1)
    await send_or_edit(callback, "🤖 <b>Выберите модуль для вопроса ИИ</b>" if modules else "Сначала создайте модуль.", kb)


@router.callback_query(F.data.startswith("module_ai_question:"))
async def module_ai_question(callback: CallbackQuery, state: FSMContext) -> None:
    module_id = int(callback.data.split(":")[1])
    if not openai_client or not OPENAI_API_KEY:
        await callback.answer("ИИ не подключен. Проверьте AI_API_KEY, AI_BASE_URL и AI_MODEL.", show_alert=True)
        return
    await state.update_data(module_id=module_id)
    await state.set_state(ModuleAIQuestion.question)
    await callback.message.edit_text(
        "🤖 Напишите вопрос по этому модулю.\n\nНапример: кто выступает 23 июля после обеда, кому нужны цветы, какой сетап нужен завтра?",
        reply_markup=kb_back(module_id).as_markup(),
    )
    await callback.answer()


@router.message(ModuleAIQuestion.question)
async def module_ai_question_answer(message: Message, state: FSMContext) -> None:
    data = await state.get_data()
    module_id = int(data["module_id"])
    question = message.text.strip()
    await message.answer("Ищу по модулю через ИИ...")
    async with SessionLocal() as session:
        owner_ids = await visible_owner_ids(session, message.from_user.id, message.from_user.username)
        module = await session.get(Module, module_id)
        if not module or module.owner_user_id not in owner_ids:
            await state.clear()
            await message.answer("Модуль не найден или доступа нет.", reply_markup=kb_main().as_markup())
            return
        context = await build_module_ai_context(session, module_id)
    try:
        answer = await ai_text_create(
            "Ты помощник менеджера мероприятия. Отвечай коротко и точно только по данным модуля. "
            "Если данных нет, прямо скажи, что в модуле это не указано. Не выдумывай.",
            f"Данные модуля:\n{context[:45000]}\n\nВопрос пользователя:\n{question}",
        )
    except Exception as exc:
        await state.clear()
        await message.answer(f"ИИ не смог ответить: {html_escape(exc)}", reply_markup=kb_back(module_id).as_markup())
        return
    await state.clear()
    for chunk in split_long_text(answer or "ИИ не вернул ответ."):
        await message.answer(html_escape(chunk))
    await message.answer("Готово.", reply_markup=kb_back(module_id).as_markup())


@router.callback_query(F.data.startswith("module_delete_ask:"))
async def module_delete_ask(callback: CallbackQuery) -> None:
    module_id = int(callback.data.split(":")[1])
    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Да, удалить", callback_data=f"module_delete:{module_id}")
    kb.button(text="❌ Отмена", callback_data=f"module:{module_id}")
    kb.adjust(1)
    await send_or_edit(callback, "Удалить модуль? Это также удалит расписание, привязки, документы и напоминания.", kb)


@router.callback_query(F.data.startswith("module_delete:"))
async def module_delete(callback: CallbackQuery) -> None:
    module_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        module = await session.get(Module, module_id)
        if module and module.owner_user_id == callback.from_user.id:
            await log_action(session, callback.from_user.id, "delete", "module", module_id, old={"name": module.name})
            await session.delete(module)
            await session.commit()
    await modules_list(callback)


@router.callback_query(F.data.startswith("module_edit:"))
async def module_edit(callback: CallbackQuery) -> None:
    module_id = int(callback.data.split(":")[1])
    kb = InlineKeyboardBuilder()
    for field, title in [("name", "Название"), ("dates", "Даты"), ("location", "Локация"), ("description", "Описание")]:
        kb.button(text=title, callback_data=f"module_edit_field:{module_id}:{field}")
    kb.button(text="⬅️ Назад", callback_data=f"module:{module_id}")
    kb.adjust(2, 2, 1)
    await send_or_edit(callback, "Что изменить в модуле?", kb)


@router.callback_query(F.data.startswith("module_edit_field:"))
async def module_edit_field(callback: CallbackQuery, state: FSMContext) -> None:
    _, module_id, field = callback.data.split(":")
    await state.update_data(module_id=int(module_id), field=field)
    await state.set_state(ModuleEdit.field)
    await callback.message.edit_text("Введите новое значение:", reply_markup=kb_back(int(module_id)).as_markup())
    await callback.answer()


@router.message(ModuleEdit.field)
async def module_edit_save(message: Message, state: FSMContext) -> None:
    data = await state.get_data()
    async with SessionLocal() as session:
        module = await session.get(Module, data["module_id"])
        if module:
            old = {"name": module.name, "dates": module.dates_label, "location": module.location, "description": module.description}
            field = data["field"]
            if field == "dates":
                start, end, label = parse_date_range(message.text)
                module.start_date, module.end_date, module.dates_label = start, end, label
            elif field == "name":
                module.name = message.text.strip()
            elif field == "location":
                module.location = message.text.strip()
            else:
                module.description = message.text.strip()
            await log_action(session, message.from_user.id, "update", "module", module.id, old=old, new={field: message.text.strip()})
            await session.commit()
            text, kb = await module_card(session, module.id)
        else:
            text, kb = "Модуль не найден.", kb_main()
    await state.clear()
    await message.answer(text, reply_markup=kb.as_markup())


@router.callback_query(F.data.startswith("speaker_create:"))
async def speaker_create(callback: CallbackQuery, state: FSMContext) -> None:
    module_id = int(callback.data.split(":")[1])
    await state.update_data(module_id=module_id)
    await state.set_state(SpeakerCreate.full_name)
    await callback.message.edit_text("👤 Введите ФИО спикера:", reply_markup=kb_back(module_id).as_markup())
    await callback.answer()


@router.message(SpeakerCreate.full_name)
async def speaker_full_name(message: Message, state: FSMContext) -> None:
    await state.update_data(full_name=message.text.strip())
    await state.set_state(SpeakerCreate.gender)
    kb = InlineKeyboardBuilder()
    kb.button(text="Мужчина", callback_data="gender:мужчина")
    kb.button(text="Женщина", callback_data="gender:женщина")
    kb.adjust(2)
    await message.answer("Пол:", reply_markup=kb.as_markup())


@router.callback_query(SpeakerCreate.gender, F.data.startswith("gender:"))
async def speaker_gender(callback: CallbackQuery, state: FSMContext) -> None:
    await state.update_data(gender=callback.data.split(":")[1])
    await state.set_state(SpeakerCreate.phone)
    await callback.message.edit_text("📞 Телефон:")
    await callback.answer()


@router.message(SpeakerCreate.phone)
async def speaker_phone(message: Message, state: FSMContext) -> None:
    await state.update_data(phone=message.text.strip())
    await state.set_state(SpeakerCreate.telegram)
    await message.answer("💬 Telegram:")


@router.message(SpeakerCreate.telegram)
async def speaker_tg(message: Message, state: FSMContext) -> None:
    await state.update_data(telegram=message.text.strip())
    await state.set_state(SpeakerCreate.email)
    await message.answer("📧 Email:")


@router.message(SpeakerCreate.email)
async def speaker_email(message: Message, state: FSMContext) -> None:
    await state.update_data(email=message.text.strip())
    await state.set_state(SpeakerCreate.organization)
    await message.answer("🏢 Организация:")


@router.message(SpeakerCreate.organization)
async def speaker_org(message: Message, state: FSMContext) -> None:
    await state.update_data(organization=message.text.strip())
    await state.set_state(SpeakerCreate.position)
    await message.answer("💼 Должность:")


@router.message(SpeakerCreate.position)
async def speaker_position(message: Message, state: FSMContext) -> None:
    await state.update_data(position=message.text.strip())
    await state.set_state(SpeakerCreate.topic)
    await message.answer("🎤 Тема выступления:")


@router.message(SpeakerCreate.topic)
async def speaker_topic(message: Message, state: FSMContext) -> None:
    await state.update_data(topic=message.text.strip())
    await state.set_state(SpeakerCreate.is_paid)
    kb = InlineKeyboardBuilder()
    kb.button(text="💰 Платный", callback_data="paid:1")
    kb.button(text="Бесплатный", callback_data="paid:0")
    kb.adjust(2)
    await message.answer("Спикер платный?", reply_markup=kb.as_markup())


@router.callback_query(SpeakerCreate.is_paid, F.data.startswith("paid:"))
async def speaker_paid(callback: CallbackQuery, state: FSMContext) -> None:
    is_paid = callback.data.endswith(":1")
    await state.update_data(is_paid=is_paid)
    if is_paid:
        await state.set_state(SpeakerCreate.amount)
        await callback.message.edit_text("💵 Укажите сумму договора:")
    else:
        await state.update_data(amount=0, contract_type="")
        await state.set_state(SpeakerCreate.setup)
        await callback.message.edit_text(f"⚙️ Укажите сетап через запятую.\nПодсказка: {SETUP_HINT}")
    await callback.answer()


@router.message(SpeakerCreate.amount)
async def speaker_amount(message: Message, state: FSMContext) -> None:
    amount = float(re.sub(r"[^\d.]", "", message.text.replace(",", ".")) or 0)
    await state.update_data(amount=amount)
    await state.set_state(SpeakerCreate.contract_type)
    kb = InlineKeyboardBuilder()
    for ctype in CONTRACT_TYPES:
        kb.button(text=ctype, callback_data=f"contract:{ctype}")
    kb.adjust(2)
    await message.answer("📄 Тип оформления:", reply_markup=kb.as_markup())


@router.callback_query(SpeakerCreate.contract_type, F.data.startswith("contract:"))
async def speaker_contract(callback: CallbackQuery, state: FSMContext) -> None:
    await state.update_data(contract_type=callback.data.split(":", 1)[1])
    await state.set_state(SpeakerCreate.setup)
    await callback.message.edit_text(f"⚙️ Укажите сетап через запятую.\nПодсказка: {SETUP_HINT}")
    await callback.answer()


@router.message(SpeakerCreate.setup)
async def speaker_setup(message: Message, state: FSMContext) -> None:
    await state.update_data(setup=split_setup(message.text))
    await state.set_state(SpeakerCreate.comment)
    await message.answer("Комментарий / особые требования или «-»:")


@router.message(SpeakerCreate.comment)
async def speaker_finish(message: Message, state: FSMContext) -> None:
    data = await state.get_data()
    async with SessionLocal() as session:
        speaker = Speaker(
            full_name=data["full_name"],
            gender=data["gender"],
            phone=data.get("phone", ""),
            telegram=data.get("telegram", ""),
            email=data.get("email", ""),
            organization=data.get("organization", ""),
            position=data.get("position", ""),
            comment="" if message.text.strip() == "-" else message.text.strip(),
            owner_user_id=message.from_user.id,
        )
        session.add(speaker)
        await session.flush()
        module = await session.get(Module, data["module_id"])
        ms = await ensure_module_speaker(
            session,
            module,
            speaker,
            topic=data.get("topic", ""),
            setup=data.get("setup", []),
            is_paid=data.get("is_paid", False),
            amount=data.get("amount", 0),
            contract_type=data.get("contract_type", ""),
        )
        await log_action(session, message.from_user.id, "create", "speaker", speaker.id, new={"full_name": speaker.full_name})
        await session.commit()
        await recreate_reminders_for_module(session, module.id, message.from_user.id)
        await session.commit()
        text, kb = await speaker_card(session, ms.id)
    await state.clear()
    await message.answer(text, reply_markup=kb.as_markup())


@router.callback_query(F.data == "speakers")
async def speakers_all(callback: CallbackQuery) -> None:
    async with SessionLocal() as session:
        owner_ids = await visible_owner_ids(session, callback.from_user.id, callback.from_user.username)
        links = (
            await session.scalars(
                select(ModuleSpeaker)
                .join(Speaker)
                .join(Module)
                .where(Module.owner_user_id.in_(owner_ids))
                .options(selectinload(ModuleSpeaker.speaker), selectinload(ModuleSpeaker.module))
                .order_by(Speaker.full_name)
            )
        ).all()
        await session.commit()
    kb = InlineKeyboardBuilder()
    for ms in links:
        kb.button(text=f"{ms.speaker.full_name} | {ms.module.name}", callback_data=f"speaker:{ms.id}")
    if links:
        kb.button(text="🗑 Удалить всех моих спикеров", callback_data="speakers_delete_all_ask")
    kb.button(text="🏠 Главное меню", callback_data="menu")
    kb.adjust(1)
    await send_or_edit(callback, "👤 <b>Спикеры</b>" if links else "Спикеров пока нет.", kb)


@router.callback_query(F.data.startswith("module_speakers:"))
async def module_speakers(callback: CallbackQuery) -> None:
    module_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        links = (
            await session.scalars(
                select(ModuleSpeaker)
                .where(ModuleSpeaker.module_id == module_id)
                .options(selectinload(ModuleSpeaker.speaker))
                .order_by(ModuleSpeaker.id)
            )
        ).all()
    kb = InlineKeyboardBuilder()
    for ms in links:
        kb.button(text=ms.speaker.full_name, callback_data=f"speaker:{ms.id}")
    kb.button(text="➕ Добавить спикера", callback_data=f"speaker_create:{module_id}")
    if links:
        kb.button(text="🗑 Удалить всех спикеров модуля", callback_data=f"module_speakers_delete_all_ask:{module_id}")
    kb.button(text="⬅️ К модулю", callback_data=f"module:{module_id}")
    kb.adjust(1)
    await send_or_edit(callback, "👤 <b>Спикеры модуля</b>" if links else "В модуле пока нет спикеров.", kb)


@router.callback_query(F.data == "speakers_delete_all_ask")
async def speakers_delete_all_ask(callback: CallbackQuery) -> None:
    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Да, удалить всех", callback_data="speakers_delete_all")
    kb.button(text="❌ Отмена", callback_data="speakers")
    kb.adjust(1)
    await send_or_edit(callback, "Удалить всех ваших спикеров из всех модулей? Расписание останется, но слоты будут без привязки к спикерам.", kb)


@router.callback_query(F.data == "speakers_delete_all")
async def speakers_delete_all(callback: CallbackQuery) -> None:
    async with SessionLocal() as session:
        speakers = (await session.scalars(select(Speaker).where(Speaker.owner_user_id == callback.from_user.id))).all()
        speaker_ids = [speaker.id for speaker in speakers]
        if speaker_ids:
            items = (await session.scalars(select(ScheduleItem).where(ScheduleItem.speaker_id.in_(speaker_ids)))).all()
            reminders = (await session.scalars(select(Reminder).where(Reminder.speaker_id.in_(speaker_ids)))).all()
            for item in items:
                item.speaker_id = None
            for reminder in reminders:
                reminder.speaker_id = None
            for speaker in speakers:
                await session.delete(speaker)
        await log_action(session, callback.from_user.id, "delete_all", "speaker", None, old={"speakers": len(speakers)})
        await session.commit()
    await send_or_edit(callback, f"✅ Удалено спикеров: {len(speaker_ids)}.", kb_main())


@router.callback_query(F.data.startswith("module_speakers_delete_all_ask:"))
async def module_speakers_delete_all_ask(callback: CallbackQuery) -> None:
    module_id = int(callback.data.split(":")[1])
    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Да, удалить всех из модуля", callback_data=f"module_speakers_delete_all:{module_id}")
    kb.button(text="❌ Отмена", callback_data=f"module_speakers:{module_id}")
    kb.adjust(1)
    await send_or_edit(callback, "Удалить всех спикеров из этого модуля? Расписание останется, но слоты будут без спикеров.", kb)


@router.callback_query(F.data.startswith("module_speakers_delete_all:"))
async def module_speakers_delete_all(callback: CallbackQuery) -> None:
    module_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        module = await session.get(Module, module_id)
        if not module or module.owner_user_id != callback.from_user.id:
            await callback.answer("Удалять всех спикеров может только владелец модуля.", show_alert=True)
            return
        links = (
            await session.scalars(
                select(ModuleSpeaker)
                .where(ModuleSpeaker.module_id == module_id)
                .options(selectinload(ModuleSpeaker.speaker))
            )
        ).all()
        speaker_ids = [link.speaker_id for link in links]
        items = (await session.scalars(select(ScheduleItem).where(ScheduleItem.module_id == module_id, ScheduleItem.speaker_id.in_(speaker_ids)))).all() if speaker_ids else []
        reminders = (await session.scalars(select(Reminder).where(Reminder.module_id == module_id, Reminder.speaker_id.in_(speaker_ids)))).all() if speaker_ids else []
        for item in items:
            item.speaker_id = None
        for reminder in reminders:
            reminder.speaker_id = None
        for link in links:
            await session.delete(link)
        await session.flush()
        for speaker_id in speaker_ids:
            still_used = await session.scalar(select(func.count(ModuleSpeaker.id)).where(ModuleSpeaker.speaker_id == speaker_id))
            speaker = await session.get(Speaker, speaker_id)
            if speaker and speaker.owner_user_id == callback.from_user.id and not still_used:
                await session.delete(speaker)
        await log_action(session, callback.from_user.id, "delete_all_module_speakers", "module", module_id, old={"speakers": len(links)})
        await session.commit()
        text, kb = await module_card(session, module_id)
    await send_or_edit(callback, f"✅ Удалено спикеров из модуля: {len(links)}.\n\n{text}", kb)


@router.callback_query(F.data.startswith("speaker:"))
async def speaker_open(callback: CallbackQuery) -> None:
    ms_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        text, kb = await speaker_card(session, ms_id)
    await send_or_edit(callback, text, kb)


@router.callback_query(F.data.startswith("speaker_delete_ask:"))
async def speaker_delete_ask(callback: CallbackQuery) -> None:
    ms_id = int(callback.data.split(":")[1])
    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Да, удалить", callback_data=f"speaker_delete:{ms_id}")
    kb.button(text="❌ Отмена", callback_data=f"speaker:{ms_id}")
    kb.adjust(1)
    await send_or_edit(callback, "Удалить спикера из модуля? Бот запросил подтверждение перед важным удалением.", kb)


@router.callback_query(F.data.startswith("speaker_delete:"))
async def speaker_delete(callback: CallbackQuery) -> None:
    ms_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        ms = await session.scalar(select(ModuleSpeaker).where(ModuleSpeaker.id == ms_id).options(selectinload(ModuleSpeaker.speaker)))
        module_id = ms.module_id if ms else None
        if ms:
            schedule_items = (
                await session.scalars(
                    select(ScheduleItem).where(
                        ScheduleItem.module_id == ms.module_id,
                        ScheduleItem.speaker_id == ms.speaker_id,
                    )
                )
            ).all()
            for item in schedule_items:
                item.speaker_id = None
            await log_action(session, callback.from_user.id, "delete", "module_speaker", ms_id, old={"speaker": ms.speaker.full_name})
            await session.delete(ms)
            await session.commit()
    if module_id:
        await module_open_fake(callback, module_id)
    else:
        await send_or_edit(callback, "Главное меню", kb_main())


async def module_open_fake(callback: CallbackQuery, module_id: int) -> None:
    async with SessionLocal() as session:
        text, kb = await module_card(session, module_id)
    await send_or_edit(callback, text, kb)


@router.callback_query(F.data.startswith("speaker_edit_setup:"))
async def speaker_edit_setup(callback: CallbackQuery, state: FSMContext) -> None:
    ms_id = int(callback.data.split(":")[1])
    await state.update_data(ms_id=ms_id)
    await state.set_state(SpeakerEdit.setup)
    await callback.message.edit_text(f"Введите новый сетап через запятую.\nПодсказка: {SETUP_HINT}")
    await callback.answer()


@router.message(SpeakerEdit.setup)
async def speaker_edit_setup_save(message: Message, state: FSMContext) -> None:
    data = await state.get_data()
    async with SessionLocal() as session:
        ms = await session.get(ModuleSpeaker, data["ms_id"])
        old = {"setup": ms.setup}
        ms.setup = split_setup(message.text)
        ms.flower_required = "цветы" in ms.setup
        await log_action(session, message.from_user.id, "update_setup", "module_speaker", ms.id, old=old, new={"setup": ms.setup})
        await session.commit()
        text, kb = await speaker_card(session, ms.id)
    await state.clear()
    await message.answer(text, reply_markup=kb.as_markup())


@router.callback_query(F.data.startswith("speaker_edit_amount:"))
async def speaker_edit_amount(callback: CallbackQuery, state: FSMContext) -> None:
    ms_id = int(callback.data.split(":")[1])
    await state.update_data(ms_id=ms_id)
    await state.set_state(SpeakerEdit.amount)
    await callback.message.edit_text("Введите новую сумму:")
    await callback.answer()


@router.message(SpeakerEdit.amount)
async def speaker_edit_amount_save(message: Message, state: FSMContext) -> None:
    data = await state.get_data()
    amount = float(re.sub(r"[^\d.]", "", message.text.replace(",", ".")) or 0)
    async with SessionLocal() as session:
        ms = await session.scalar(select(ModuleSpeaker).where(ModuleSpeaker.id == data["ms_id"]).options(selectinload(ModuleSpeaker.document_status)))
        old = {"amount": ms.amount}
        ms.amount = amount
        ms.is_paid = amount > 0
        if ms.is_paid and not ms.document_status:
            session.add(DocumentStatus(module_speaker_id=ms.id))
        await log_action(session, message.from_user.id, "update_amount", "module_speaker", ms.id, old=old, new={"amount": amount})
        await session.commit()
        text, kb = await speaker_card(session, ms.id)
    await state.clear()
    await message.answer(text, reply_markup=kb.as_markup())


@router.callback_query(F.data.startswith("speaker_edit_data:"))
async def speaker_edit_data(callback: CallbackQuery) -> None:
    ms_id = int(callback.data.split(":")[1])
    kb = InlineKeyboardBuilder()
    kb.button(text="ФИО", callback_data=f"speaker_field:{ms_id}:full_name")
    kb.button(text="Телефон", callback_data=f"speaker_field:{ms_id}:phone")
    kb.button(text="Telegram", callback_data=f"speaker_field:{ms_id}:telegram")
    kb.button(text="Email", callback_data=f"speaker_field:{ms_id}:email")
    kb.button(text="Тема", callback_data=f"speaker_field:{ms_id}:topic")
    kb.button(text="Комментарий", callback_data=f"speaker_field:{ms_id}:comment")
    kb.button(text="⬅️ Назад", callback_data=f"speaker:{ms_id}")
    kb.adjust(2, 2, 2, 1)
    await send_or_edit(callback, "Что изменить?", kb)


@router.callback_query(F.data.startswith("speaker_field:"))
async def speaker_edit_field(callback: CallbackQuery, state: FSMContext) -> None:
    _, ms_id, field = callback.data.split(":")
    await state.update_data(ms_id=int(ms_id), field=field)
    await state.set_state(SpeakerEdit.data)
    await callback.message.edit_text("Введите новое значение:")
    await callback.answer()


@router.message(SpeakerEdit.data)
async def speaker_edit_data_save(message: Message, state: FSMContext) -> None:
    data = await state.get_data()
    async with SessionLocal() as session:
        ms = await session.scalar(select(ModuleSpeaker).where(ModuleSpeaker.id == data["ms_id"]).options(selectinload(ModuleSpeaker.speaker)))
        field = data["field"]
        target = ms if field == "topic" else ms.speaker
        old = {field: getattr(target, field)}
        setattr(target, field, message.text.strip())
        await log_action(session, message.from_user.id, "update", "speaker", ms.speaker_id, old=old, new={field: message.text.strip()})
        await session.commit()
        text, kb = await speaker_card(session, ms.id)
    await state.clear()
    await message.answer(text, reply_markup=kb.as_markup())


@router.callback_query(F.data.startswith("docs:"))
async def docs_open(callback: CallbackQuery) -> None:
    ms_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        ms = await session.scalar(
            select(ModuleSpeaker)
            .where(ModuleSpeaker.id == ms_id)
            .options(selectinload(ModuleSpeaker.speaker), selectinload(ModuleSpeaker.document_status))
        )
        if not ms:
            await send_or_edit(callback, "Спикер не найден.", kb_main())
            return
        if not ms.document_status:
            ms.is_paid = True
            ms.document_status = DocumentStatus(module_speaker_id=ms.id)
            await session.commit()
        progress = doc_progress(ms.document_status)
        await session.commit()
        lines = [f"📄 <b>Документы: {html_escape(ms.speaker.full_name)}</b>", ""]
        for field, title in DOC_FIELDS:
            lines.append(f"{'✅' if getattr(ms.document_status, field) else '❌'} {html_escape(title)}")
        lines.append(f"\n<b>Прогресс:</b> {progress}%")
    kb = InlineKeyboardBuilder()
    for field, title in DOC_FIELDS:
        kb.button(text=f"Переключить: {title}", callback_data=f"doc_toggle:{ms_id}:{field}")
    kb.button(text="⬅️ К спикеру", callback_data=f"speaker:{ms_id}")
    kb.adjust(1)
    await send_or_edit(callback, "\n".join(lines), kb)


@router.callback_query(F.data.startswith("doc_toggle:"))
async def doc_toggle(callback: CallbackQuery) -> None:
    _, ms_id, field = callback.data.split(":")
    async with SessionLocal() as session:
        ms = await session.scalar(select(ModuleSpeaker).where(ModuleSpeaker.id == int(ms_id)).options(selectinload(ModuleSpeaker.document_status)))
        if ms and ms.document_status and field in dict(DOC_FIELDS):
            old = {field: getattr(ms.document_status, field)}
            setattr(ms.document_status, field, not getattr(ms.document_status, field))
            doc_progress(ms.document_status)
            await log_action(session, callback.from_user.id, "toggle_document", "document_status", ms.document_status.id, old=old, new={field: getattr(ms.document_status, field)})
            await session.commit()
    await docs_open(callback)


@router.callback_query(F.data.startswith("module_payments:") | (F.data == "payments_all"))
async def payments(callback: CallbackQuery) -> None:
    module_id = None if callback.data == "payments_all" else int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        owner_ids = await visible_owner_ids(session, callback.from_user.id, callback.from_user.username)
        stmt = (
            select(ModuleSpeaker)
            .join(Speaker)
            .join(Module)
            .where(Module.owner_user_id.in_(owner_ids), ModuleSpeaker.is_paid == True)
            .options(selectinload(ModuleSpeaker.speaker), selectinload(ModuleSpeaker.module), selectinload(ModuleSpeaker.document_status))
        )
        if module_id:
            stmt = stmt.where(ModuleSpeaker.module_id == module_id)
        links = (await session.scalars(stmt)).all()
    total = sum(x.amount for x in links)
    avg = round(sum(doc_progress(x.document_status) for x in links) / len(links)) if links else 0
    title = "💰 <b>Оплаты и документы</b>"
    if module_id and links:
        title += f"\n📁 Модуль: {html_escape(links[0].module.name)}"
    lines = [title, "", f"Платных спикеров: {len(links)}", f"Общая сумма: {money(total)}", f"Средний прогресс документов: {avg}%", ""]
    kb = InlineKeyboardBuilder()
    for i, ms in enumerate(links, 1):
        lines += [
            f"{i}. <b>{html_escape(ms.speaker.full_name)}</b>",
            f"Сумма: {money(ms.amount)}",
            f"Тип: {html_escape(ms.contract_type or 'не указан')}",
            f"Прогресс: {doc_progress(ms.document_status)}%",
            "",
        ]
        kb.button(text=f"📄 {ms.speaker.full_name}", callback_data=f"docs:{ms.id}")
    kb.button(text="🏠 Главное меню", callback_data="menu")
    kb.adjust(1)
    await send_or_edit(callback, "\n".join(lines) if links else "Платных спикеров пока нет.", kb)


@router.callback_query(F.data.startswith("module_flowers:") | (F.data == "flowers_all"))
async def flowers(callback: CallbackQuery) -> None:
    module_id = None if callback.data == "flowers_all" else int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        owner_ids = await visible_owner_ids(session, callback.from_user.id, callback.from_user.username)
        stmt = (
            select(ModuleSpeaker)
            .join(Speaker)
            .join(Module)
            .where(Module.owner_user_id.in_(owner_ids))
            .options(selectinload(ModuleSpeaker.speaker), selectinload(ModuleSpeaker.module))
        )
        if module_id:
            stmt = stmt.where(ModuleSpeaker.module_id == module_id)
        links = [x for x in (await session.scalars(stmt)).all() if x.flower_required or x.speaker.gender.lower().startswith("жен")]
        rows = []
        for ms in links:
            item = await session.scalar(
                select(ScheduleItem)
                .where(ScheduleItem.module_id == ms.module_id, ScheduleItem.speaker_id == ms.speaker_id)
                .order_by(ScheduleItem.date, ScheduleItem.start_time)
            )
            rows.append((ms, item))
    module_name = rows[0][0].module.name if module_id and rows else "всем модулям"
    lines = [
        f"🌸 <b>Цветы: {html_escape(module_name)}</b>",
        "",
        f"👤 Женщин-спикеров: <b>{len(rows)}</b>",
        f"💐 Букетов к заказу: <b>{len(rows)}</b>",
        "",
    ]
    for i, (ms, item) in enumerate(rows, 1):
        when = f"{item.date:%d.%m.%Y} | {time_label(item.start_time)}-{time_label(item.end_time)}" if item else "время не указано"
        topic = ms.topic or (item.title if item else "")
        lines += [
            f"<b>{i}. {html_escape(ms.speaker.full_name)}</b>",
            f"📁 {html_escape(ms.module.name)}",
            f"⏰ {html_escape(when)}",
            f"🎤 {html_escape(topic or 'тема не указана')}",
            "",
        ]
    if not rows:
        lines.append("Пока нет женщин-спикеров или отмеченных букетов.")
    await send_or_edit(callback, "\n".join(lines), kb_back(module_id))


@router.callback_query(F.data.startswith("upload_schedule:"))
async def upload_schedule(callback: CallbackQuery, state: FSMContext) -> None:
    module_id = int(callback.data.split(":")[1])
    await state.update_data(module_id=module_id)
    await state.set_state(ScheduleUpload.waiting_file)
    await callback.message.edit_text(
        "📎 Загрузите файл расписания: xlsx, xls, csv, txt, docx, pdf, jpg или png.",
        reply_markup=kb_back(module_id).as_markup(),
    )
    await callback.answer()


async def writable_owner_id(session: AsyncSession, user_id: int, username: str | None) -> int:
    owners = await visible_owner_ids(session, user_id, username)
    for owner_id in owners:
        if owner_id != user_id:
            return owner_id
    return user_id


@router.callback_query(F.data == "pretty_schedule")
async def pretty_schedule_menu(callback: CallbackQuery) -> None:
    async with SessionLocal() as session:
        owner_ids = await visible_owner_ids(session, callback.from_user.id, callback.from_user.username)
        images = (
            await session.scalars(
                select(PrettyScheduleImage)
                .where(PrettyScheduleImage.owner_user_id.in_(owner_ids))
                .order_by(PrettyScheduleImage.schedule_date, PrettyScheduleImage.created_at)
            )
        ).all()
    kb = InlineKeyboardBuilder()
    for image in images:
        kb.button(text=image.day_label, callback_data=f"pretty_schedule_show:{image.id}")
    kb.button(text="➕ Загрузить картинку дня", callback_data="pretty_schedule_upload")
    kb.button(text="🏠 Главное меню", callback_data="menu")
    kb.adjust(1)
    text = "🖼 <b>Красивое расписание</b>\n\nВыберите день, чтобы открыть картинку." if images else "🖼 <b>Красивое расписание</b>\n\nПока нет загруженных картинок."
    await send_or_edit(callback, text, kb)


@router.callback_query(F.data == "pretty_schedule_upload")
async def pretty_schedule_upload(callback: CallbackQuery, state: FSMContext) -> None:
    await state.set_state(PrettyScheduleUpload.day_label)
    await callback.message.edit_text("Введите день для красивого расписания. Например: 23 июля или 23.07.2026")
    await callback.answer()


@router.message(PrettyScheduleUpload.day_label)
async def pretty_schedule_day_save(message: Message, state: FSMContext) -> None:
    label = message.text.strip()
    if not label:
        await message.answer("Напишите день текстом, например: 23 июля.")
        return
    parsed_date = parse_date_value(label)
    await state.update_data(day_label=label, schedule_date=parsed_date.isoformat() if parsed_date else None)
    await state.set_state(PrettyScheduleUpload.image)
    await message.answer("Теперь отправьте картинку расписания для этого дня.")


@router.message(PrettyScheduleUpload.image, F.photo | F.document)
async def pretty_schedule_image_save(message: Message, state: FSMContext, bot: Bot) -> None:
    data = await state.get_data()
    if message.photo:
        photo: PhotoSize = message.photo[-1]
        src_name = f"pretty_{photo.file_id}.jpg"
        file = await bot.get_file(photo.file_id)
    elif message.document:
        doc: Document = message.document
        src_name = doc.file_name or f"pretty_{doc.file_id}"
        if not re.search(r"\.(jpg|jpeg|png|webp)$", src_name, flags=re.I):
            await message.answer("Нужна картинка: jpg, png или webp.")
            return
        file = await bot.get_file(doc.file_id)
    else:
        await message.answer("Отправьте картинку.")
        return

    safe_name = re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", src_name)
    dst = UPLOAD_DIR / f"pretty_{datetime.utcnow():%Y%m%d%H%M%S}_{safe_name}"
    await bot.download_file(file.file_path, destination=dst)
    async with SessionLocal() as session:
        owner_id = await writable_owner_id(session, message.from_user.id, message.from_user.username)
        image = PrettyScheduleImage(
            owner_user_id=owner_id,
            day_label=data["day_label"],
            schedule_date=date.fromisoformat(data["schedule_date"]) if data.get("schedule_date") else None,
            file_name=src_name,
            file_path=str(dst),
        )
        session.add(image)
        await log_action(session, message.from_user.id, "upload_pretty_schedule", "pretty_schedule", None, new={"day": data["day_label"]})
        await session.commit()
    await state.clear()
    permissions = await effective_permissions(message.from_user.id, message.from_user.username)
    await message.answer("✅ Красивая картинка расписания сохранена.", reply_markup=kb_main(permissions).as_markup())


@router.callback_query(F.data.startswith("pretty_schedule_show:"))
async def pretty_schedule_show(callback: CallbackQuery) -> None:
    image_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        owner_ids = await visible_owner_ids(session, callback.from_user.id, callback.from_user.username)
        image = await session.get(PrettyScheduleImage, image_id)
        if not image or image.owner_user_id not in owner_ids:
            await callback.answer("Картинка не найдена.", show_alert=True)
            return
    path = Path(image.file_path)
    if not path.exists():
        await callback.answer("Файл картинки не найден на сервере.", show_alert=True)
        return
    await callback.message.answer_photo(FSInputFile(path), caption=f"🖼 {html_escape(image.day_label)}", reply_markup=kb_back().as_markup())
    await callback.answer()


@router.message(ScheduleUpload.waiting_file, F.document | F.photo)
async def receive_schedule_file(message: Message, state: FSMContext, bot: Bot) -> None:
    data = await state.get_data()
    module_id = data["module_id"]
    src_name = "schedule.jpg"
    if message.document:
        doc: Document = message.document
        src_name = doc.file_name or f"file_{doc.file_id}"
        file = await bot.get_file(doc.file_id)
    else:
        photo: PhotoSize = message.photo[-1]
        src_name = f"photo_{photo.file_id}.jpg"
        file = await bot.get_file(photo.file_id)
    safe_name = re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", src_name)
    dst = UPLOAD_DIR / f"{datetime.utcnow():%Y%m%d%H%M%S}_{safe_name}"
    await bot.download_file(file.file_path, destination=dst)

    await message.answer("Файл получен. Извлекаю текст и привожу расписание к единому JSON-формату.")
    try:
        extracted = extract_file_text(dst)
        async with SessionLocal() as session:
            module = await session.get(Module, module_id)
            settings = await get_settings(session, message.from_user.id)
            if "CELL_BLOCKS:" in extracted:
                parsed = parse_excel_cell_blocks_to_schedule(module, extracted)
                if settings.ai_enabled and openai_client and OPENAI_API_KEY:
                    parsed = await enrich_fixed_excel_schedule_with_ai(module, parsed, extracted, require_ai=False)
            else:
                if settings.ai_enabled and openai_client and OPENAI_API_KEY:
                    parsed = await parse_schedule_smart_with_ai(module, extracted, require_ai=True)
                    if len(split_layout_matrix_by_day(extracted)) <= 1:
                        parsed = await normalize_schedule_with_ai(module, parsed, extracted, require_ai=False)
                else:
                    parsed = heuristic_schedule(module, extracted)
                    has_local_items = any(day.get("items") for day in parsed.get("days", []))
                    if not has_local_items:
                        raise RuntimeError("Локальный разбор не нашел слоты расписания. Включите ИИ или загрузите XLSX/PDF с видимой сеткой дат и времени.")
            validate_schedule_coverage(parsed, extracted)
            try:
                validate_schedule_quality(parsed)
            except ScheduleQualityError as quality_exc:
                if settings.ai_enabled and openai_client and OPENAI_API_KEY:
                    await message.answer("Kimi склеил несколько событий в один длинный слот. Отправляю на повторную разбивку и проверку...")
                    parsed = await repair_schedule_with_ai(module, parsed, extracted, quality_exc)
                    validate_schedule_coverage(parsed, extracted)
                    validate_schedule_quality(parsed)
                else:
                    raise
            upload = UploadedFile(
                module_id=module_id,
                file_name=src_name,
                file_type=dst.suffix.lower().lstrip("."),
                file_path=str(dst),
                parsed_json=parsed,
            )
            session.add(upload)
            await session.flush()
            await session.commit()
            preview = schedule_preview(parsed)
    except Exception as exc:
        await state.clear()
        await message.answer(
            "❌ Расписание не сохранено.\n\n"
            f"{html_escape(exc)}\n\n"
            "Если ИИ включен - проверьте AI_API_KEY/AI_BASE_URL/AI_MODEL. Если ИИ выключен - загрузите XLSX/PDF с понятной сеткой дат, времени и ячеек.",
            reply_markup=kb_back(module_id).as_markup(),
        )
        return
    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Сохранить расписание", callback_data=f"schedule_confirm:{upload.id}")
    kb.button(text="❌ Отмена", callback_data=f"module:{module_id}")
    kb.adjust(1)
    await state.clear()
    await message.answer(preview, reply_markup=kb.as_markup())


def extract_file_text(path: Path) -> str:
    ext = path.suffix.lower()
    try:
        if ext in {".csv"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if ext in {".txt"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if ext == ".xlsx" and openpyxl is not None:
            return extract_excel_text_with_layout(path)
        if ext in {".xlsx", ".xls"} and pd is not None:
            frames = pd.read_excel(path, sheet_name=None, header=None)
            return "\n\n".join(f"Лист: {name}\n{df.fillna('').to_csv(index=False, header=False)}" for name, df in frames.items())
        if ext == ".docx" and docx is not None:
            document = docx.Document(str(path))
            parts = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        if ext == ".pdf" and pdfplumber is not None:
            return extract_pdf_text_with_layout(path)
        if ext in {".jpg", ".jpeg", ".png"} and pytesseract is not None and Image is not None:
            return pytesseract.image_to_string(Image.open(path), lang="rus+eng")
    except Exception as exc:
        log.exception("file extract failed: %s", exc)
        return f"Не удалось извлечь файл автоматически: {exc}"
    return "Не удалось извлечь текст: установите нужную библиотеку или OCR. Можно отправить расписание текстом."


def extract_excel_text_with_layout(path: Path) -> str:
    wb = openpyxl.load_workbook(path, data_only=True)
    parts = [
        "=== EXCEL_LAYOUT_RULES ===",
        "This is an Excel schedule. Use CELL_BLOCKS/MERGED_CELLS as the main source, not CSV text.",
        "A merged cell is one schedule item. Its start_time is the first time row it covers, and end_time is the end of the last time row it covers.",
        "If one cell contains multiple lines, they are title/speaker/comment lines inside the same slot, not separate slots.",
        "Example: a cell containing 'ЛПР\\nЖога' across 11:30-13:00 is one item 11:30-13:00 with title 'ЛПР Жога'.",
        "Never split one Excel cell into several 15-minute items.",
        "",
    ]
    for ws in wb.worksheets:
        parts.append(f"=== SHEET: {ws.title} ===")
        merged_lookup = merged_cell_lookup(ws)
        date_headers = detect_excel_date_headers(ws, merged_lookup)
        time_rows = detect_excel_time_rows(ws)
        merged_lines = describe_excel_merged_ranges(ws)
        if merged_lines:
            parts.append("MERGED_RANGES:")
            parts.extend(merged_lines)
        if date_headers and time_rows:
            parts.append("DATE_COLUMNS: " + " | ".join(f"{col}:{value}" for col, value in date_headers.items()))
            parts.append("TIME_ROWS: " + " | ".join(f"{row}:{start}-{end}" for row, start, end in time_rows))
            parts.append("CELL_BLOCKS:")
            parts.extend(build_excel_cell_blocks(ws, date_headers, time_rows, merged_lookup))
            parts.append("ALL_NON_EMPTY_CELLS:")
            parts.extend(describe_excel_non_empty_cells(ws, merged_lookup))
        else:
            parts.append("RAW_CELLS:")
            parts.extend(describe_excel_non_empty_cells(ws, merged_lookup))
        parts.append("")
    return "\n".join(parts)


def clean_cell_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, time):
        return value.strftime("%H:%M")
    if isinstance(value, datetime):
        if value.date() == date(1899, 12, 30) or value.date() == date(1900, 1, 1):
            return value.time().strftime("%H:%M")
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, (int, float)) and 0 <= float(value) < 1:
        total_minutes = round(float(value) * 24 * 60)
        return f"{total_minutes // 60:02d}:{total_minutes % 60:02d}"
    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    return " / ".join(line for line in lines if line)


def describe_excel_merged_ranges(ws: Any) -> list[str]:
    lines = []
    for merged_range in ws.merged_cells.ranges:
        value = clean_cell_value(ws.cell(merged_range.min_row, merged_range.min_col).value)
        if value:
            lines.append(f"{merged_range.coord} | rows {merged_range.min_row}-{merged_range.max_row} | cols {merged_range.min_col}-{merged_range.max_col} | {value}")
    return lines


def describe_excel_non_empty_cells(ws: Any, merged_lookup: dict[tuple[int, int], Any]) -> list[str]:
    lines = []
    seen = set()
    for row in range(1, ws.max_row + 1):
        for col in range(1, ws.max_column + 1):
            merged_range = merged_lookup.get((row, col))
            if merged_range:
                key = (merged_range.min_row, merged_range.min_col, merged_range.max_row, merged_range.max_col)
                if key in seen:
                    continue
                seen.add(key)
                value = clean_cell_value(ws.cell(merged_range.min_row, merged_range.min_col).value)
                address = merged_range.coord
            else:
                value = clean_cell_value(ws.cell(row, col).value)
                address = f"{get_column_letter(col)}{row}" if get_column_letter else f"R{row}C{col}"
            if value:
                lines.append(f"{address} | row {row} | col {col} | {value}")
    return lines


def merged_cell_lookup(ws: Any) -> dict[tuple[int, int], Any]:
    lookup = {}
    for merged_range in ws.merged_cells.ranges:
        for row in range(merged_range.min_row, merged_range.max_row + 1):
            for col in range(merged_range.min_col, merged_range.max_col + 1):
                lookup[(row, col)] = merged_range
    return lookup


def cell_value_with_merge(ws: Any, row: int, col: int, merged_lookup: dict[tuple[int, int], Any]) -> str:
    merged_range = merged_lookup.get((row, col))
    if merged_range:
        return clean_cell_value(ws.cell(merged_range.min_row, merged_range.min_col).value)
    return clean_cell_value(ws.cell(row, col).value)


def detect_excel_date_headers(ws: Any, merged_lookup: dict[tuple[int, int], Any]) -> dict[int, str]:
    headers = {}
    max_scan_rows = min(ws.max_row, 30)
    for row in range(1, max_scan_rows + 1):
        for col in range(1, ws.max_column + 1):
            value = cell_value_with_merge(ws, row, col, merged_lookup)
            parsed = parse_date_value(value)
            if parsed:
                merged_range = merged_lookup.get((row, col))
                min_col = merged_range.min_col if merged_range else col
                max_col = merged_range.max_col if merged_range else col
                for c in range(min_col, max_col + 1):
                    headers[c] = parsed.isoformat()
    return headers


def detect_excel_time_rows(ws: Any) -> list[tuple[int, str, str]]:
    rows = []
    single_time_rows = []
    for row in range(1, ws.max_row + 1):
        row_start = None
        row_end = None
        row_times = []
        for col in range(1, ws.max_column + 1):
            value = clean_cell_value(ws.cell(row, col).value)
            tm = re.search(r"(\d{1,2}[:.]\d{2})(?::\d{2})?\s*[-–]\s*(\d{1,2}[:.]\d{2})(?::\d{2})?", value)
            if tm:
                row_start = tm.group(1).replace(".", ":")
                row_end = tm.group(2).replace(".", ":")
                break
            for single in re.findall(r"\b\d{1,2}[:.]\d{2}(?::\d{2})?\b", value):
                single = ":".join(single.replace(".", ":").split(":")[:2])
                row_times.append(single.replace(".", ":"))
        if not row_start and len(row_times) >= 2:
            row_start, row_end = row_times[0], row_times[1]
        if row_start and row_end:
            rows.append((row, row_start, row_end))
        elif row_times:
            single_time_rows.append((row, row_times[0]))
    if rows:
        return rows
    for (row, start), (_next_row, end) in zip(single_time_rows, single_time_rows[1:]):
        if time_to_minutes(end) > time_to_minutes(start):
            rows.append((row, start, end))
    return rows


def build_excel_cell_blocks(
    ws: Any,
    date_headers: dict[int, str],
    time_rows: list[tuple[int, str, str]],
    merged_lookup: dict[tuple[int, int], Any],
) -> list[str]:
    row_times = {row: (start, end) for row, start, end in time_rows}
    time_row_numbers = [row for row, _, _ in time_rows]
    blocks = []
    seen = set()
    for row in time_row_numbers:
        for col, day in sorted(date_headers.items()):
            value = cell_value_with_merge(ws, row, col, merged_lookup)
            if not value:
                continue
            merged_range = merged_lookup.get((row, col))
            min_row = merged_range.min_row if merged_range else row
            max_row = merged_range.max_row if merged_range else row
            min_col = merged_range.min_col if merged_range else col
            max_col = merged_range.max_col if merged_range else col
            key = (min_row, max_row, min_col, max_col)
            if key in seen:
                continue
            seen.add(key)
            covered_rows = [r for r in time_row_numbers if min_row <= r <= max_row]
            if not covered_rows:
                covered_rows = [row]
            start = row_times[covered_rows[0]][0]
            end = row_times[covered_rows[-1]][1]
            address = f"{get_column_letter(min_col)}{min_row}:{get_column_letter(max_col)}{max_row}" if get_column_letter else f"R{min_row}C{min_col}:R{max_row}C{max_col}"
            blocks.append(f"{day} | {start}-{end} | {address} | {value}")
    return blocks


def extract_pdf_text_with_layout(path: Path) -> str:
    """Extract PDF text plus a coordinate-based schedule matrix.

    Many event grids are visually clear but text extraction interleaves days:
    one left time column, several date columns, and merged cells that span
    multiple 15-minute rows. This layout-aware representation gives the AI the
    row/column context it needs to recover exact dates and time ranges.
    """
    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page_index, page in enumerate(pdf.pages, 1):
            raw_text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
            parts.append(f"\n=== PAGE {page_index} RAW_TEXT ===\n{raw_text}")

            words = page.extract_words(
                x_tolerance=1,
                y_tolerance=2,
                keep_blank_chars=False,
                use_text_flow=False,
            )
            matrix = build_schedule_matrix_from_words(words, page.width, page.lines)
            if matrix:
                parts.append(f"\n=== PAGE {page_index} LAYOUT_MATRIX ===")
                parts.append(matrix)

            tables = page.extract_tables() or []
            if tables:
                parts.append(f"\n=== PAGE {page_index} TABLES ===")
                for table in tables:
                    for row in table:
                        parts.append(" | ".join((cell or "").strip() for cell in row))
    return "\n".join(parts)


def build_schedule_matrix_from_words(words: list[dict[str, Any]], page_width: float, page_lines: list[dict[str, Any]] | None = None) -> str:
    date_re = re.compile(r"\b\d{1,2}\.\d{1,2}\.\d{4}\b")
    time_re = re.compile(r"^\d{1,2}:\d{2}\s*[-–]\s*\d{1,2}:\d{2}$")

    date_headers = []
    for word in words:
        if date_re.fullmatch(word["text"]):
            date_headers.append(
                {
                    "date": word["text"],
                    "x0": float(word["x0"]),
                    "x1": float(word["x1"]),
                    "center": (float(word["x0"]) + float(word["x1"])) / 2,
                    "top": float(word["top"]),
                }
            )
    date_headers = sorted(date_headers, key=lambda x: x["center"])
    if len(date_headers) < 2:
        return ""

    time_words = []
    first_date_x = date_headers[0]["x0"]
    for word in words:
        text = normalize_extracted_time(word["text"])
        if time_re.fullmatch(text) and float(word["x1"]) < first_date_x:
            start, end = [x.strip() for x in re.split(r"[-–]", text, maxsplit=1)]
            time_words.append(
                {
                    "start": start,
                    "end": end,
                    "x1": float(word["x1"]),
                    "top": float(word["top"]),
                    "bottom": float(word["bottom"]),
                    "center": (float(word["top"]) + float(word["bottom"])) / 2,
                }
            )
    time_words = sorted(time_words, key=lambda x: x["center"])
    if len(time_words) < 2:
        return ""

    time_col_right = max(t["x1"] for t in time_words)
    column_ranges = []
    for i, header in enumerate(date_headers):
        left = (
            time_col_right + 2
            if i == 0
            else (date_headers[i - 1]["center"] + header["center"]) / 2
        )
        right = (
            page_width
            if i == len(date_headers) - 1
            else (header["center"] + date_headers[i + 1]["center"]) / 2
        )
        column_ranges.append((header["date"], left, right))

    row_ranges = []
    for i, slot in enumerate(time_words):
        top = (
            slot["top"] - 2
            if i == 0
            else (time_words[i - 1]["center"] + slot["center"]) / 2
        )
        bottom = (
            slot["bottom"] + 2
            if i == len(time_words) - 1
            else (slot["center"] + time_words[i + 1]["center"]) / 2
        )
        row_ranges.append((slot["start"], slot["end"], top, bottom))

    lines = [
        "The grid below was reconstructed from PDF coordinates.",
        "Columns are dates. Rows are time slots.",
        "In this program grid, schedule items go one after another without gaps inside a day.",
        "The end_time of an item is normally the start_time of the next real item in the same date column.",
        "Rows inside one visual activity block are details of that same item, not separate agenda items.",
        "Do not split a large activity into 15-minute fragments.",
        "",
        "DATES: " + " | ".join(d for d, _, _ in column_ranges),
    ]

    for start, end, top, bottom in row_ranges:
        cells = []
        for day, left, right in column_ranges:
            cell_words = [
                word
                for word in words
                if top <= float(word["bottom"]) < bottom
                and left <= ((float(word["x0"]) + float(word["x1"])) / 2) < right
                and not date_re.fullmatch(word["text"])
            ]
            text = join_positioned_words(cell_words)
            cells.append(f"{day}: {text}" if text else f"{day}:")
        lines.append(f"{start}-{end} || " + " || ".join(cells))

    visual_blocks = build_visual_block_hints(column_ranges, row_ranges, page_lines or [])
    if visual_blocks:
        lines.append("")
        lines.append("VISUAL_BLOCKS:")
        lines.extend(visual_blocks)

    return "\n".join(lines)


def build_visual_block_hints(
    column_ranges: list[tuple[str, float, float]],
    row_ranges: list[tuple[str, str, float, float]],
    page_lines: list[dict[str, Any]],
) -> list[str]:
    if not page_lines:
        return []

    def boundary_time(y: float) -> str | None:
        nearest = None
        nearest_dist = 999.0
        for start, _end, top, _bottom in row_ranges:
            dist = abs(y - top)
            if dist < nearest_dist:
                nearest = start
                nearest_dist = dist
        return nearest if nearest_dist <= 3.0 else None

    hints = []
    seen = set()
    for day, left, right in column_ranges:
        boundaries = []
        for line in page_lines:
            if abs(float(line.get("height", 0))) > 0.01:
                continue
            x0 = float(line["x0"])
            x1 = float(line["x1"])
            y = float(line["top"])
            if x0 <= left + 8 and x1 >= right - 8:
                t = boundary_time(y)
                if t:
                    boundaries.append((y, t))
        boundaries = sorted(set(boundaries))
        for (_y1, start), (_y2, end) in zip(boundaries, boundaries[1:]):
            if start == end:
                continue
            key = (day, start, end)
            if key in seen:
                continue
            seen.add(key)
            if time_to_minutes(end) - time_to_minutes(start) >= 30:
                hints.append(f"{day}: {start}-{end}")
    return hints


def normalize_extracted_time(text: str) -> str:
    return (
        text.strip()
        .replace("–", "-")
        .replace("—", "-")
        .replace(".", ":")
        .replace(" ", "")
    )


def join_positioned_words(words: list[dict[str, Any]]) -> str:
    if not words:
        return ""
    lines: dict[int, list[dict[str, Any]]] = defaultdict(list)
    for word in words:
        lines[round(float(word["top"]) / 3)].append(word)

    chunks = []
    for _, line_words in sorted(lines.items()):
        row = sorted(line_words, key=lambda w: float(w["x0"]))
        chunks.append(" ".join(w["text"] for w in row))
    text = " / ".join(chunk for chunk in chunks if chunk.strip())
    return re.sub(r"\s+", " ", text).strip()


def schedule_json_schema() -> dict[str, Any]:
    return {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "module_name": {"type": "string"},
            "days": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "date": {"type": "string"},
                        "items": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "start_time": {"type": "string"},
                                    "end_time": {"type": "string"},
                                    "title": {"type": "string"},
                                    "speaker_name": {"type": "string"},
                                    "speaker_names": {"type": "array", "items": {"type": "string"}},
                                    "location": {"type": "string"},
                                    "format": {"type": "string"},
                                    "comment": {"type": "string"},
                                },
                                "required": ["start_time", "end_time", "title", "speaker_name", "speaker_names", "location", "format", "comment"],
                            },
                        },
                    },
                    "required": ["date", "items"],
                },
            },
        },
        "required": ["module_name", "days"],
    }


async def parse_schedule_with_ai(module: Module, text: str, require_ai: bool = False) -> dict[str, Any]:
    if openai_client and OPENAI_API_KEY:
        schema = {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "module_name": {"type": "string"},
                "days": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "date": {"type": "string"},
                            "items": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "additionalProperties": False,
                                    "properties": {
                                        "start_time": {"type": "string"},
                                        "end_time": {"type": "string"},
                                        "title": {"type": "string"},
                                        "speaker_name": {"type": "string"},
                                        "speaker_names": {"type": "array", "items": {"type": "string"}},
                                        "location": {"type": "string"},
                                        "format": {"type": "string"},
                                        "comment": {"type": "string"},
                                    },
                                    "required": ["start_time", "end_time", "title", "speaker_name", "speaker_names", "location", "format", "comment"],
                                },
                            },
                        },
                        "required": ["date", "items"],
                    },
                },
            },
            "required": ["module_name", "days"],
        }
        try:
            response = await openai_client.responses.create(
                model=OPENAI_MODEL,
                input=[
                    {
                        "role": "system",
                        "content": (
                            "Ты профессионально извлекаешь расписания мероприятий из PDF/Excel/табличных сеток. "
                            "Верни только валидный JSON по схеме. Даты приводи к YYYY-MM-DD. "
                            "КРИТИЧЕСКИ ВАЖНО: не склеивай весь день или большой фрагмент дня в один слот. "
                            "Если внутри текста подряд встречаются маркеры новых событий ('Лекция', 'Мастер-класс', 'Обед', 'Ужин', 'Кофе-брейк', 'Перерыв', 'Открытие дня', 'Тема:'), "
                            "создавай отдельные items. Конец одного события равен началу следующего события. "
                            "Если один текстовый блок содержит много разных спикеров и много разных тем, это не один слот, а несколько последовательных событий. "
                            "Разделяй такие блоки по словам 'Лекция', 'Мастер-класс', 'Кофе-брейк', 'Обед', 'Перерыв', а также по смене ведущего/спикера. "
                            "Слот 9:30-22:30 почти всегда ошибка, если внутри перечислены разные лекции, мастер-классы, обед, кофе-брейки или разные ведущие. "
                            "Имена людей отделяй от описаний: после запятой обычно идет должность/роль, а не продолжение имени. "
                            "Если в тексте есть LAYOUT_MATRIX, считай ее главным источником: колонки - это даты, строки - временные слоты. "
                            "Текст в одной колонке даты относится именно к этой дате. "
                            "Обязательно восстанавливай длительность событий по соседним строкам одной колонки: "
                            "если название, тема или спикер явно продолжаются текстом в следующих 15-минутных строках, "
                            "объедини их в один item с start_time первой строки и end_time последней строки. "
                            "В этой сетке слоты идут подряд без дыр внутри дня: конец события обычно равен началу следующего настоящего события в той же колонке. "
                            "Не режь крупный блок на 15-минутные куски. Внутренние строки блока считай описанием/частями этого же item. "
                            "Если в исходном тексте есть EXCEL_LAYOUT_RULES или CELL_BLOCKS, считай каждый CELL_BLOCK одной Excel-ячейкой и одним слотом. "
                            "Время слота бери из CELL_BLOCKS. Не дели текст одной ячейки на отдельные события. "
                            "Например CELL_BLOCK 'ЛПР Жога' 11:30-13:00 - это один item 11:30-13:00, а не два item. "
                            "Например, если 'Командообразование на улице по станциям' включает строки '(4 станции по командам', 'Стрельба...', 'Урок связи...', 'Квест...', "
                            "сохрани это одним item с началом первого блока и концом перед следующим крупным событием. "
                            "Если появляется явно новая крупная активность после блока, заверши предыдущий item временем начала новой активности. "
                            "Например, если в 10:00-10:15 указано 'Знакомство', а в 10:15-10:30 '(или мосты)', "
                            "создай один item 10:00-10:30 с title 'Знакомство (или мосты)'. "
                            "Если в строке 7:45-8:00 в колонке 24.07 стоит 'Пробежка с Владимиром Волошиным', "
                            "создай item на 2026-07-24 с start_time 07:45 и end_time 08:00. "
                            "Отделяй спикеров от темы: speaker_names должен содержать только людей, а не активности. "
                            "Например 'Светлана Попова-Смойлик', 'Сергей Леваненков', 'Владимир Волошин' - это спикеры; "
                            "'Завтрак', 'Обед', 'Кофе-брейк', 'Открытие дня', 'Знакомство', 'Жога', 'ЛПР' - не спикеры. "
                            "Если спикеров несколько, заполни speaker_names всеми именами, а speaker_name продублируй строкой через '; '. "
                            "Если спикера нет, speaker_name='' и speaker_names=[]."
                        ),
                    },
                    {"role": "user", "content": f"Модуль: {module.name}\nДаты: {module.dates_label}\nТекст файла:\n{text[:50000]}"},
                ],
                text={"format": {"type": "json_schema", "name": "schedule", "schema": schema, "strict": True}},
            )
            return json.loads(response.output_text)
        except Exception as exc:
            log.exception("OpenAI schedule parse failed: %s", exc)
            if require_ai:
                raise RuntimeError(f"ИИ не смог распознать расписание: {exc}") from exc
    if require_ai:
        raise RuntimeError("OPENAI_API_KEY не задан или OpenAI API недоступен. Расписание не сохранено без ИИ-проверки.")
    return heuristic_schedule(module, text)


def split_layout_matrix_by_day(source_text: str) -> dict[str, str]:
    day_lines: dict[str, list[str]] = defaultdict(list)
    for line in source_text.splitlines():
        if " || " not in line:
            continue
        time_part, cells_part = line.split(" || ", 1)
        tm = re.fullmatch(r"(\d{1,2}:\d{2})-(\d{1,2}:\d{2})", time_part.strip())
        if not tm:
            continue
        for raw_cell in cells_part.split(" || "):
            if ":" not in raw_cell:
                continue
            day_raw, cell_text = raw_cell.split(":", 1)
            day = parse_date_value(day_raw.strip())
            if not day:
                continue
            day_key = day.isoformat()
            text_value = cell_text.strip()
            if text_value:
                day_lines[day_key].append(f"{tm.group(1)}-{tm.group(2)} | {text_value}")
            else:
                day_lines[day_key].append(f"{tm.group(1)}-{tm.group(2)} |")
    result = {}
    for day, lines in day_lines.items():
        result[day] = "\n".join(
            [
                "=== LAYOUT_MATRIX_ONE_DAY ===",
                f"DATE: {day}",
                "Rows are exact time slots for this date. Split events when the activity, format or speaker changes.",
                "If a row is a continuation of the previous row, merge it into the previous item.",
                "",
                *lines,
            ]
        )
    return result


async def parse_schedule_smart_with_ai(module: Module, source_text: str, require_ai: bool = False) -> dict[str, Any]:
    day_sections = split_layout_matrix_by_day(source_text)
    if len(day_sections) <= 1:
        return await parse_schedule_with_ai(module, source_text, require_ai=require_ai)

    parsed_days = []
    for day, day_text in sorted(day_sections.items()):
        day_parsed = await parse_schedule_with_ai(module, day_text, require_ai=require_ai)
        for parsed_day in day_parsed.get("days", []):
            if not parsed_day.get("date"):
                parsed_day["date"] = day
            parsed_days.append(parsed_day)
    return {"module_name": module.name, "days": parsed_days}


async def normalize_schedule_with_ai(module: Module, parsed: dict[str, Any], source_text: str, require_ai: bool = False) -> dict[str, Any]:
    if not openai_client or not OPENAI_API_KEY:
        if require_ai:
            raise RuntimeError("OPENAI_API_KEY не задан. Нельзя сохранить расписание без ИИ-проверки.")
        return parsed

    schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "module_name": {"type": "string"},
            "days": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "date": {"type": "string"},
                        "items": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "start_time": {"type": "string"},
                                    "end_time": {"type": "string"},
                                    "title": {"type": "string"},
                                    "speaker_name": {"type": "string"},
                                    "speaker_names": {"type": "array", "items": {"type": "string"}},
                                    "location": {"type": "string"},
                                    "format": {"type": "string"},
                                    "comment": {"type": "string"},
                                },
                                "required": ["start_time", "end_time", "title", "speaker_name", "speaker_names", "location", "format", "comment"],
                            },
                        },
                    },
                    "required": ["date", "items"],
                },
            },
        },
        "required": ["module_name", "days"],
    }
    try:
        response = await openai_client.responses.create(
            model=OPENAI_MODEL,
            input=[
                {
                    "role": "system",
                    "content": (
                        "Ты проверяешь уже распознанное расписание перед сохранением в базу. "
                        "Исправь фрагментацию: строки, которые являются продолжением темы, описанием эксперта или должностью спикера, должны быть внутри одного item. "
                        "Пример: 'Мастер-класс' + '«Техники построения систем»' + 'Михаил Федоренко, эксперт по развитию' + "
                        "'целостной личности, построению' + 'устойчивых систем и эффективных команд' - это один item, "
                        "title='Мастер-класс «Техники построения систем»', speaker_names=['Михаил Федоренко'], comment содержит описание эксперта. "
                        "Не создавай отдельные слоты из фраз вроде 'целостной личности, построению' или 'устойчивых систем и эффективных команд'. "
                        "Если исходный текст содержит CELL_BLOCKS, каждый CELL_BLOCK является атомарным слотом Excel: не разбивай его внутренние строки на несколько событий. "
                        "Например 'ЛПР Жога' в одной ячейке 11:30-13:00 - это один item 11:30-13:00. "
                        "Если VISUAL_BLOCKS есть в исходном тексте, используй их как границы крупных блоков. "
                        "Слоты в одном дне должны идти последовательно: конец одного обычно равен началу следующего. "
                        "Верни только JSON по схеме."
                    ),
                },
                {
                    "role": "user",
                    "content": (
                        f"Модуль: {module.name}\n"
                        f"Черновой JSON:\n{json.dumps(parsed, ensure_ascii=False)[:30000]}\n\n"
                        f"Исходная матрица/текст для проверки:\n{source_text[:30000]}"
                    ),
                },
            ],
            text={"format": {"type": "json_schema", "name": "normalized_schedule", "schema": schema, "strict": True}},
        )
        return json.loads(response.output_text)
    except Exception as exc:
        log.exception("OpenAI schedule normalize failed: %s", exc)
        if require_ai:
            raise RuntimeError(f"ИИ-проверка расписания не прошла: {exc}") from exc
        return parsed


async def repair_schedule_with_ai(
    module: Module,
    parsed: dict[str, Any],
    source_text: str,
    quality_error: Exception,
) -> dict[str, Any]:
    if not openai_client or not OPENAI_API_KEY:
        raise quality_error
    day_match = re.search(r"(\d{4}-\d{2}-\d{2})", str(quality_error))
    repair_source = source_text
    if day_match:
        repair_source = split_layout_matrix_by_day(source_text).get(day_match.group(1), source_text)
    system_prompt = (
        "Ты исправляешь плохо распознанное расписание мероприятия. "
        "Главная ошибка: модель склеила несколько разных событий в один длинный слот. "
        "Твоя задача - вернуть полный JSON расписания, но разбить длинные склеенные слоты на отдельные реальные события. "
        "Используй исходный текст/матрицу как главный источник времени. "
        "Если есть LAYOUT_MATRIX: строки - это временные интервалы, колонки - даты. "
        "Новая активность начинается, когда в тексте появляются маркеры: Лекция, Мастер-класс, Обед, Ужин, Кофе-брейк, Перерыв, Открытие дня, Сессия, Практикум, Панель, Рефлексия, Дорога. "
        "Описания должностей после имени спикера оставляй в comment, а не делай отдельным слотом. "
        "Не делай слот длиннее 4 часов, если внутри есть несколько разных тем, несколько разных форматов или много спикеров. "
        "Если точное время конца не написано, конец события равен началу следующего события в этом же дне. "
        "Спикеры - только имена людей. Активности, еда, дорога и названия организаций не являются спикерами. "
        "Верни только валидный JSON по схеме."
    )
    user_prompt = (
        f"Модуль: {module.name}\n"
        f"Даты модуля: {module.dates_label}\n"
        f"Ошибка качества: {quality_error}\n\n"
        f"Плохой JSON, который надо исправить:\n{json.dumps(parsed, ensure_ascii=False)[:30000]}\n\n"
        f"Исходный текст/матрица файла:\n{repair_source[:50000]}"
    )
    repaired = await ai_json_create(system_prompt, user_prompt, schedule_json_schema(), "repaired_schedule")
    if day_match:
        target_day = day_match.group(1)
        repaired_days = [day for day in repaired.get("days", []) if day.get("date") == target_day]
        if repaired_days:
            merged = dict(parsed)
            merged["days"] = [
                repaired_days[0] if day.get("date") == target_day else day
                for day in parsed.get("days", [])
            ]
            return merged
    return repaired


def validate_schedule_coverage(parsed: dict[str, Any], source_text: str) -> None:
    parsed_items = sum(len(day.get("items", [])) for day in parsed.get("days", []))
    if parsed_items == 0:
        raise RuntimeError("ИИ не вернул ни одного слота расписания. Расписание не сохранено.")

    has_structured_source = "CELL_BLOCKS:" in source_text or "LAYOUT_MATRIX" in source_text or "VISUAL_BLOCKS:" in source_text
    if has_structured_source and parsed_items < 3:
        raise RuntimeError(
            f"ИИ вернул слишком мало слотов: {parsed_items}. "
            "Расписание не сохранено, чтобы не потерять данные. Повторите загрузку или проверьте модель."
        )


def validate_schedule_quality(parsed: dict[str, Any]) -> None:
    event_markers = (
        "лекция", "мастер-класс", "обед", "ужин", "кофе", "перерыв", "модератор",
        "ведущий", "тренер", "делегирование", "управление", "открытие", "сессия",
        "практикум", "панель", "встречает", "завтрак",
    )
    for day in parsed.get("days", []):
        for item in day.get("items", []):
            start = time_to_minutes(item.get("start_time"))
            end = time_to_minutes(item.get("end_time"))
            duration = end - start if start and end else 0
            text_blob = " ".join(
                str(item.get(key, ""))
                for key in ("title", "speaker_name", "comment", "format", "location")
            ).lower()
            marker_count = sum(1 for marker in event_markers if marker in text_blob)
            too_many_speakers = len(item_speaker_names(item)) >= 6
            if duration >= 300 and (marker_count >= 3 or too_many_speakers):
                raise ScheduleQualityError(
                    "ИИ склеил несколько событий в один слишком длинный слот. "
                    f"Проблемный слот: {day.get('date')} {item.get('start_time')}-{item.get('end_time')}. "
                    "Запущу повторную разбивку этого фрагмента через ИИ."
                )


def parse_excel_cell_blocks_to_schedule(module: Module, source_text: str) -> dict[str, Any]:
    days: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for line in source_text.splitlines():
        match = re.match(r"^(\d{4}-\d{2}-\d{2})\s+\|\s+(\d{1,2}:\d{2})-(\d{1,2}:\d{2})\s+\|\s+([^|]+)\|\s*(.+)$", line)
        if not match:
            continue
        day, start, end, address, raw_text = match.groups()
        raw_text = raw_text.strip()
        if not raw_text:
            continue
        speakers = extract_speaker_candidates(raw_text)
        title = excel_block_title(raw_text, speakers)
        item = {
            "start_time": start,
            "end_time": end,
            "title": title,
            "speaker_name": "; ".join(speakers),
            "speaker_names": speakers,
            "location": "",
            "format": "",
            "comment": f"excel_cell={address.strip()}; raw={raw_text}",
        }
        days[day].append(item)
    for items in days.values():
        merge_offline_day_items(items)
    return {
        "module_name": module.name,
        "days": [
            {
                "date": day,
                "items": sorted(items, key=lambda item: (time_to_minutes(item["start_time"]), time_to_minutes(item["end_time"]))),
            }
            for day, items in sorted(days.items())
        ],
    }


def merge_offline_day_items(items: list[dict[str, Any]]) -> None:
    if not items:
        return
    items.sort(key=lambda item: (time_to_minutes(item.get("start_time")), time_to_minutes(item.get("end_time"))))
    merged: list[dict[str, Any]] = []
    for item in items:
        if not merged:
            merged.append(item)
            continue
        prev = merged[-1]
        contiguous = prev.get("end_time") == item.get("start_time")
        same_title = clean_title_key(prev.get("title", "")) == clean_title_key(item.get("title", ""))
        continuation = is_activity_continuation(prev.get("title", ""), item.get("title", "")) or item.get("title", "")[:1].islower()
        if contiguous and (same_title or continuation):
            if not same_title and item.get("title") and item["title"] not in prev.get("title", ""):
                prev["title"] = f"{prev.get('title', '')} {item['title']}".strip()
            speakers = unique_names((prev.get("speaker_names") or []) + (item.get("speaker_names") or []))
            prev["speaker_names"] = speakers
            prev["speaker_name"] = "; ".join(speakers)
            prev["end_time"] = item.get("end_time") or prev.get("end_time")
            if item.get("comment"):
                prev["comment"] = f"{prev.get('comment', '')}; {item['comment']}".strip("; ")
        else:
            merged.append(item)
    items[:] = merged


def clean_title_key(value: str) -> str:
    return re.sub(r"\W+", "", (value or "").lower())


def excel_block_title(raw_text: str, speakers: Iterable[str]) -> str:
    parts = [part.strip(" -–—,.;:") for part in re.split(r"\s*/\s*|\n+", raw_text) if part.strip(" -–—,.;:")]
    cleaned_parts = [remove_speaker_names_from_title(part, speakers) for part in parts]
    cleaned_parts = [part for part in cleaned_parts if part and clean_cell_title_part(part)]
    if cleaned_parts:
        return " ".join(cleaned_parts)
    return remove_speaker_names_from_title(raw_text, speakers)


def clean_cell_title_part(text: str) -> bool:
    low = text.lower().strip()
    if not low:
        return False
    service_prefixes = ("excel_cell=", "raw=")
    return not low.startswith(service_prefixes)


async def enrich_fixed_excel_schedule_with_ai(module: Module, parsed: dict[str, Any], source_text: str, require_ai: bool = False) -> dict[str, Any]:
    if not openai_client or not OPENAI_API_KEY:
        if require_ai:
            raise RuntimeError("OPENAI_API_KEY не задан. Нельзя сохранить Excel-расписание без ИИ-проверки.")
        return parsed

    expected_keys = fixed_schedule_keys(parsed)
    schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "module_name": {"type": "string"},
            "days": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "date": {"type": "string"},
                        "items": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "start_time": {"type": "string"},
                                    "end_time": {"type": "string"},
                                    "title": {"type": "string"},
                                    "speaker_name": {"type": "string"},
                                    "speaker_names": {"type": "array", "items": {"type": "string"}},
                                    "location": {"type": "string"},
                                    "format": {"type": "string"},
                                    "comment": {"type": "string"},
                                },
                                "required": ["start_time", "end_time", "title", "speaker_name", "speaker_names", "location", "format", "comment"],
                            },
                        },
                    },
                    "required": ["date", "items"],
                },
            },
        },
        "required": ["module_name", "days"],
    }
    try:
        response = await openai_client.responses.create(
            model=OPENAI_MODEL,
            input=[
                {
                    "role": "system",
                    "content": (
                        "Ты чистишь Excel-расписание, но структура уже зафиксирована кодом. "
                        "Строго запрещено менять date, start_time, end_time и количество items. "
                        "Один входной item = одна Excel-ячейка = один выходной item. "
                        "Твоя задача только: аккуратно выделить title, speaker_names, format, location, comment. "
                        "Если в ячейке написано 'ЛПР / Жога', это один item с title='ЛПР Жога'. "
                        "Не дели ячейку на несколько событий. Не объединяй соседние ячейки. "
                        "Спикерами считаются только имена людей. Активности, еда, дорога и места не являются спикерами. "
                        "Если спикера нет, speaker_name='' и speaker_names=[]. "
                        "Верни JSON по схеме с тем же набором date/start_time/end_time."
                    ),
                },
                {
                    "role": "user",
                    "content": (
                        f"Модуль: {module.name}\n"
                        f"Фиксированное расписание, сохраняй все даты и время без изменений:\n"
                        f"{json.dumps(parsed, ensure_ascii=False)[:45000]}\n\n"
                        f"Исходные Excel-блоки для справки:\n{source_text[:15000]}"
                    ),
                },
            ],
            text={"format": {"type": "json_schema", "name": "fixed_excel_schedule", "schema": schema, "strict": True}},
        )
        enriched = json.loads(response.output_text)
        if fixed_schedule_keys(enriched) != expected_keys:
            log.warning("AI changed fixed Excel schedule shape; using deterministic schedule")
            return parsed
        return enriched
    except Exception as exc:
        log.exception("OpenAI Excel enrich failed: %s", exc)
        if require_ai:
            raise RuntimeError(f"ИИ-проверка Excel-расписания не прошла: {exc}") from exc
        return parsed


def fixed_schedule_keys(parsed: dict[str, Any]) -> list[tuple[str, str, str]]:
    keys = []
    for day in parsed.get("days", []):
        day_value = day.get("date", "")
        for item in day.get("items", []):
            keys.append((day_value, item.get("start_time", ""), item.get("end_time", "")))
    return sorted(keys)


def heuristic_schedule(module: Module, text: str) -> dict[str, Any]:
    layout = heuristic_schedule_from_layout_matrix(module, text)
    if layout.get("days"):
        return layout

    days: dict[str, list[dict[str, str]]] = {}
    current_date = module.start_date or datetime.now().date()
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        found_date = parse_date_value(line)
        if found_date:
            current_date = found_date
        tm = re.search(r"(\d{1,2}[:.]\d{2})\s*[-–]\s*(\d{1,2}[:.]\d{2})", line)
        if tm:
            rest = line[tm.end():].strip(" -–|")
            parts = [x.strip() for x in re.split(r"\s+\|\s+|;", rest) if x.strip()]
            item = {
                "start_time": tm.group(1).replace(".", ":"),
                "end_time": tm.group(2).replace(".", ":"),
                "title": parts[0] if parts else rest,
                "speaker_name": parts[1] if len(parts) > 1 else "",
                "location": parts[2] if len(parts) > 2 else module.location,
                "format": parts[3] if len(parts) > 3 else "",
                "comment": "",
            }
            days.setdefault(current_date.isoformat(), []).append(item)
    return {"module_name": module.name, "days": [{"date": day, "items": items} for day, items in days.items()]}


def heuristic_schedule_from_layout_matrix(module: Module, text: str) -> dict[str, Any]:
    rows = []
    visual_blocks: dict[str, list[tuple[str, str]]] = defaultdict(list)
    for line in text.splitlines():
        block_match = re.fullmatch(r"(\d{1,2}\.\d{1,2}\.\d{4}):\s*(\d{1,2}:\d{2})-(\d{1,2}:\d{2})", line.strip())
        if block_match:
            block_day = parse_date_value(block_match.group(1))
            if block_day:
                visual_blocks[block_day.isoformat()].append((block_match.group(2), block_match.group(3)))
            continue
        if " || " not in line:
            continue
        time_part, cells_part = line.split(" || ", 1)
        tm = re.fullmatch(r"(\d{1,2}:\d{2})-(\d{1,2}:\d{2})", time_part.strip())
        if not tm:
            continue
        cells = {}
        for raw_cell in cells_part.split(" || "):
            if ":" not in raw_cell:
                continue
            day_raw, cell_text = raw_cell.split(":", 1)
            day = parse_date_value(day_raw.strip())
            if day:
                cells[day.isoformat()] = cell_text.strip()
        rows.append({"start": tm.group(1), "end": tm.group(2), "cells": cells})

    if not rows:
        return {"module_name": module.name, "days": []}

    by_day: dict[str, list[dict[str, str]]] = defaultdict(list)
    active: dict[str, dict[str, str] | None] = defaultdict(lambda: None)

    for row in rows:
        for day, cell_text in row["cells"].items():
            current = active[day]
            if not cell_text:
                continue

            should_continue = bool(
                current
                and is_activity_continuation(current.get("title", ""), cell_text)
            )
            if should_continue:
                current["title"] = f"{current['title']} {cell_text}".strip()
                current["end_time"] = row["end"]
            else:
                if current:
                    current["end_time"] = row["start"]
                speakers = extract_speaker_candidates(cell_text)
                title = remove_speaker_names_from_title(cell_text, speakers)
                item = {
                    "start_time": row["start"],
                    "end_time": row["end"],
                    "title": title,
                    "speaker_name": "; ".join(speakers),
                    "speaker_names": speakers,
                    "location": module.location or "",
                    "format": "",
                    "comment": "fallback_layout_parse",
                }
                by_day[day].append(item)
                active[day] = item

    for items in by_day.values():
        merge_offline_day_items(items)
    apply_visual_blocks_to_items(by_day, visual_blocks)

    return {
        "module_name": module.name,
        "days": [{"date": day, "items": items} for day, items in sorted(by_day.items())],
    }


def apply_visual_blocks_to_items(
    by_day: dict[str, list[dict[str, str]]],
    visual_blocks: dict[str, list[tuple[str, str]]],
) -> None:
    for day, blocks in visual_blocks.items():
        items = by_day.get(day, [])
        if not items:
            continue
        consumed: set[int] = set()
        new_items: list[dict[str, str]] = []
        for block_start, block_end in sorted(blocks, key=lambda pair: time_to_minutes(pair[0])):
            start_min = time_to_minutes(block_start)
            end_min = time_to_minutes(block_end)
            inside = [
                (idx, item)
                for idx, item in enumerate(items)
                if idx not in consumed
                and start_min <= time_to_minutes(item["start_time"]) < end_min
            ]
            if not inside:
                continue
            first_idx, first = inside[0]
            title_low = first.get("title", "").lower()
            can_use_block_start = (
                len(inside) > 1
                or "командообраз" in title_low
                or "веревочный" in title_low
                or "мастер-класс" in title_low
                or "сессия" in title_low
            )
            if can_use_block_start:
                first["start_time"] = block_start
            first["end_time"] = block_end
            for idx, extra in inside[1:]:
                if is_activity_continuation(first.get("title", ""), extra.get("title", "")):
                    first["title"] = f"{first['title']} {extra['title']}".strip()
                    first["speaker_name"] = "; ".join(unique_names([first.get("speaker_name", ""), extra.get("speaker_name", "")]))
                    first["speaker_names"] = unique_names((first.get("speaker_names") or []) + (extra.get("speaker_names") or []))
                    consumed.add(idx)
            consumed.add(first_idx)
            new_items.append(first)

        for idx, item in enumerate(items):
            if idx not in consumed:
                new_items.append(item)
        by_day[day] = sorted(new_items, key=lambda item: (time_to_minutes(item["start_time"]), time_to_minutes(item["end_time"])))

    normalize_contiguous_day_items(by_day)


def normalize_contiguous_day_items(by_day: dict[str, list[dict[str, str]]]) -> None:
    for day, items in by_day.items():
        items.sort(key=lambda item: (time_to_minutes(item["start_time"]), time_to_minutes(item["end_time"])))
        for prev, current in zip(items, items[1:]):
            prev_end = time_to_minutes(prev.get("end_time"))
            current_start = time_to_minutes(current.get("start_time"))
            current_end = time_to_minutes(current.get("end_time"))
            if prev_end and current_start and current_start < prev_end < current_end:
                current["start_time"] = prev["end_time"]
                current_start = prev_end
            if prev_end and current_start and 0 < current_start - prev_end <= 90:
                current["start_time"] = prev["end_time"]


def schedule_preview(parsed: dict[str, Any]) -> str:
    lines = [f"📅 <b>Распознанное расписание: {html_escape(parsed.get('module_name', ''))}</b>", ""]
    shown = 0
    total = sum(len(day.get("items", [])) for day in parsed.get("days", []))
    for day in parsed.get("days", []):
        lines.append(f"<b>{html_escape(day.get('date', ''))}</b>")
        for item in day.get("items", []):
            if shown >= 35:
                lines.append("")
                lines.append(f"Показаны первые {shown} слотов из {total}. Полное расписание сохранится после подтверждения.")
                text = "\n".join(lines)
                return text[:3600]
            speakers = item_speaker_names(item)
            speaker_text = ", ".join(speakers) if speakers else "Спикера нет"
            lines.append(f"{item.get('start_time', '')}-{item.get('end_time', '')} | <b>{html_escape(speaker_text)}</b>")
            lines.append(f"Тема: {html_escape(item.get('title', ''))}")
            shown += 1
        lines.append("")
    text = "\n".join(lines) or "ИИ не нашел расписание в файле."
    return text if len(text) <= 3600 else text[:3500] + "\n\n...предпросмотр сокращен, полное расписание сохранится после подтверждения."


@router.callback_query(F.data.startswith("schedule_confirm:"))
async def schedule_confirm(callback: CallbackQuery) -> None:
    upload_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        upload = await session.get(UploadedFile, upload_id)
        if not upload:
            await send_or_edit(callback, "Файл не найден.", kb_main())
            return
        await save_schedule_json(session, upload.module_id, upload.parsed_json, callback.from_user.id)
        await recreate_reminders_for_module(session, upload.module_id, callback.from_user.id)
        await log_action(session, callback.from_user.id, "upload_schedule", "module", upload.module_id, new=upload.parsed_json)
        await session.commit()
        text, kb = await module_card(session, upload.module_id)
    await send_or_edit(callback, "✅ Расписание сохранено.\n\n" + text, kb)


@router.callback_query(F.data.startswith("schedule_delete_menu:"))
async def schedule_delete_menu(callback: CallbackQuery) -> None:
    module_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        days = (
            await session.scalars(
                select(ScheduleItem.date)
                .where(ScheduleItem.module_id == module_id)
                .distinct()
                .order_by(ScheduleItem.date)
            )
        ).all()
    kb = InlineKeyboardBuilder()
    for d in days:
        kb.button(text=f"🗑 {d:%d.%m.%Y}", callback_data=f"schedule_delete_day_ask:{module_id}:{d.isoformat()}")
    kb.button(text="🗑 Удалить все расписание", callback_data=f"schedule_delete_all_ask:{module_id}")
    kb.button(text="⬅️ К модулю", callback_data=f"module:{module_id}")
    kb.adjust(1)
    await send_or_edit(callback, "Что удалить из расписания?", kb)


@router.callback_query(F.data.startswith("schedule_delete_day_ask:"))
async def schedule_delete_day_ask(callback: CallbackQuery) -> None:
    _, module_id, day = callback.data.split(":")
    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Да, удалить день", callback_data=f"schedule_delete_day:{module_id}:{day}")
    kb.button(text="❌ Отмена", callback_data=f"schedule_delete_menu:{module_id}")
    kb.adjust(1)
    await send_or_edit(callback, f"Удалить расписание за {html_escape(day)}?", kb)


@router.callback_query(F.data.startswith("schedule_delete_all_ask:"))
async def schedule_delete_all_ask(callback: CallbackQuery) -> None:
    module_id = int(callback.data.split(":")[1])
    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Да, удалить все", callback_data=f"schedule_delete_all:{module_id}")
    kb.button(text="❌ Отмена", callback_data=f"schedule_delete_menu:{module_id}")
    kb.adjust(1)
    await send_or_edit(callback, "Удалить все расписание модуля? Спикеры останутся в модуле.", kb)


@router.callback_query(F.data.startswith("schedule_delete_day:"))
async def schedule_delete_day(callback: CallbackQuery) -> None:
    _, module_id_raw, day_raw = callback.data.split(":")
    module_id = int(module_id_raw)
    selected_day = date.fromisoformat(day_raw)
    async with SessionLocal() as session:
        items = (
            await session.scalars(
                select(ScheduleItem).where(ScheduleItem.module_id == module_id, ScheduleItem.date == selected_day)
            )
        ).all()
        for item in items:
            reminders = (
                await session.scalars(select(Reminder).where(Reminder.schedule_item_id == item.id))
            ).all()
            for reminder in reminders:
                await session.delete(reminder)
            await session.delete(item)
        await log_action(session, callback.from_user.id, "delete_schedule_day", "module", module_id, old={"date": day_raw, "items": len(items)})
        await session.commit()
        text, kb = await module_card(session, module_id)
    await send_or_edit(callback, f"✅ Удалено расписание за {selected_day:%d.%m.%Y}.\n\n{text}", kb)


@router.callback_query(F.data.startswith("schedule_delete_all:"))
async def schedule_delete_all(callback: CallbackQuery) -> None:
    module_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        items = (await session.scalars(select(ScheduleItem).where(ScheduleItem.module_id == module_id))).all()
        reminders = (await session.scalars(select(Reminder).where(Reminder.module_id == module_id))).all()
        for reminder in reminders:
            await session.delete(reminder)
        for item in items:
            await session.delete(item)
        await log_action(session, callback.from_user.id, "delete_schedule_all", "module", module_id, old={"items": len(items)})
        await session.commit()
        text, kb = await module_card(session, module_id)
    await send_or_edit(callback, "✅ Расписание модуля удалено.\n\n" + text, kb)


async def save_schedule_json(session: AsyncSession, module_id: int, parsed: dict[str, Any], user_id: int) -> None:
    module = await session.get(Module, module_id)
    for old in (await session.scalars(select(ScheduleItem).where(ScheduleItem.module_id == module_id))).all():
        await session.delete(old)
    await session.flush()
    for day in parsed.get("days", []):
        item_date = parse_date_value(day.get("date", "")) or module.start_date or datetime.now().date()
        for item in day.get("items", []):
            speakers = item_speaker_names(item)
            title = item.get("title", "")
            if speakers:
                for speaker_name in speakers:
                    speaker = await find_speaker(session, user_id, speaker_name)
                    if not speaker:
                        speaker = Speaker(
                            full_name=speaker_name,
                            gender=infer_gender_from_name(speaker_name),
                            owner_user_id=user_id,
                        )
                        session.add(speaker)
                        await session.flush()
                    ms = await ensure_module_speaker(session, module, speaker, topic=title)
                    if speaker.gender.lower().startswith("жен"):
                        ms.flower_required = True
                    session.add(
                        ScheduleItem(
                            module_id=module_id,
                            speaker_id=speaker.id,
                            date=item_date,
                            start_time=parse_time(item.get("start_time", "")),
                            end_time=parse_time(item.get("end_time", "")),
                            title=title,
                            location=item.get("location", ""),
                            format=item.get("format", ""),
                            comment=item.get("comment", ""),
                        )
                    )
            else:
                session.add(
                    ScheduleItem(
                        module_id=module_id,
                        speaker_id=None,
                        date=item_date,
                        start_time=parse_time(item.get("start_time", "")),
                        end_time=parse_time(item.get("end_time", "")),
                        title=title,
                        location=item.get("location", ""),
                        format=item.get("format", ""),
                        comment=item.get("comment", ""),
                    )
                )


@router.callback_query(F.data.startswith("manual_schedule:"))
async def manual_schedule(callback: CallbackQuery, state: FSMContext) -> None:
    module_id = int(callback.data.split(":")[1])
    await state.update_data(module_id=module_id)
    await state.set_state(ManualSchedule.item_date)
    await callback.message.edit_text("📅 Дата слота: например 2026-06-24 или 24 июня")
    await callback.answer()


@router.message(ManualSchedule.item_date)
async def manual_date(message: Message, state: FSMContext) -> None:
    parsed = parse_date_value(message.text)
    if not parsed:
        await message.answer("Не понял дату. Попробуйте формат 2026-06-24.")
        return
    await state.update_data(date=parsed.isoformat())
    await state.set_state(ManualSchedule.start_time)
    await message.answer("⏰ Время начала, например 14:30:")


@router.message(ManualSchedule.start_time)
async def manual_start(message: Message, state: FSMContext) -> None:
    if not parse_time(message.text):
        await message.answer("Не понял время. Например: 14:30")
        return
    await state.update_data(start_time=message.text)
    await state.set_state(ManualSchedule.end_time)
    await message.answer("⏰ Время окончания:")


@router.message(ManualSchedule.end_time)
async def manual_end(message: Message, state: FSMContext) -> None:
    if not parse_time(message.text):
        await message.answer("Не понял время. Например: 15:20")
        return
    await state.update_data(end_time=message.text)
    await state.set_state(ManualSchedule.title)
    await message.answer("🎤 Тема / название слота:")


@router.message(ManualSchedule.title)
async def manual_title(message: Message, state: FSMContext) -> None:
    await state.update_data(title=message.text.strip())
    await state.set_state(ManualSchedule.speaker_name)
    await message.answer("👤 ФИО спикера или «-», если это общий слот:")


@router.message(ManualSchedule.speaker_name)
async def manual_speaker(message: Message, state: FSMContext) -> None:
    await state.update_data(speaker_name="" if message.text.strip() == "-" else message.text.strip())
    await state.set_state(ManualSchedule.location)
    await message.answer("📍 Зал / локация:")


@router.message(ManualSchedule.location)
async def manual_location(message: Message, state: FSMContext) -> None:
    await state.update_data(location=message.text.strip())
    await state.set_state(ManualSchedule.item_format)
    await message.answer("Формат: лекция, панель, мастер-класс или «-»:")


@router.message(ManualSchedule.item_format)
async def manual_format(message: Message, state: FSMContext) -> None:
    await state.update_data(format="" if message.text.strip() == "-" else message.text.strip())
    await state.set_state(ManualSchedule.comment)
    await message.answer("Комментарий или «-»:")


@router.message(ManualSchedule.comment)
async def manual_finish(message: Message, state: FSMContext) -> None:
    data = await state.get_data()
    async with SessionLocal() as session:
        module = await session.get(Module, data["module_id"])
        speaker = await find_speaker(session, message.from_user.id, data.get("speaker_name", ""))
        if not speaker and data.get("speaker_name"):
            speaker = Speaker(full_name=data["speaker_name"], owner_user_id=message.from_user.id)
            session.add(speaker)
            await session.flush()
        if speaker:
            await ensure_module_speaker(session, module, speaker, topic=data["title"])
        item = ScheduleItem(
            module_id=module.id,
            speaker_id=speaker.id if speaker else None,
            date=date.fromisoformat(data["date"]),
            start_time=parse_time(data["start_time"]),
            end_time=parse_time(data["end_time"]),
            title=data["title"],
            location=data["location"],
            format=data["format"],
            comment="" if message.text.strip() == "-" else message.text.strip(),
        )
        session.add(item)
        await session.flush()
        await recreate_reminders_for_module(session, module.id, message.from_user.id)
        await log_action(session, message.from_user.id, "create", "schedule_item", item.id, new={"title": item.title})
        await session.commit()
        text, kb = await module_card(session, module.id)
    await state.clear()
    await message.answer("✅ Слот добавлен.\n\n" + text, reply_markup=kb.as_markup())


@router.callback_query(F.data.startswith("schedule_days:") | (F.data == "schedule_all"))
async def schedule_days(callback: CallbackQuery) -> None:
    module_id = None if callback.data == "schedule_all" else int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        owner_ids = await visible_owner_ids(session, callback.from_user.id, callback.from_user.username)
        stmt = select(ScheduleItem).join(Module).where(Module.owner_user_id.in_(owner_ids)).order_by(ScheduleItem.date)
        if module_id:
            stmt = stmt.where(ScheduleItem.module_id == module_id)
        days = sorted({x.date for x in (await session.scalars(stmt)).all()})
    kb = InlineKeyboardBuilder()
    for d in days:
        kb.button(text=d.strftime("%d.%m.%Y"), callback_data=f"schedule_day:{module_id or 0}:{d.isoformat()}")
    if module_id:
        kb.button(text="🗑 Удалить расписание", callback_data=f"schedule_delete_menu:{module_id}")
        kb.button(text="⬅️ К модулю", callback_data=f"module:{module_id}")
    kb.button(text="🏠 Главное меню", callback_data="menu")
    kb.adjust(1)
    await send_or_edit(callback, "📅 Выберите день:" if days else "Расписание пока пустое.", kb)


@router.callback_query(F.data.startswith("schedule_day:"))
async def schedule_day(callback: CallbackQuery) -> None:
    _, module_id_raw, day_raw = callback.data.split(":")
    module_id = int(module_id_raw)
    selected_day = date.fromisoformat(day_raw)
    async with SessionLocal() as session:
        owner_ids = await visible_owner_ids(session, callback.from_user.id, callback.from_user.username)
        stmt = (
            select(ScheduleItem)
            .join(Module)
            .where(Module.owner_user_id.in_(owner_ids), ScheduleItem.date == selected_day)
            .options(selectinload(ScheduleItem.speaker), selectinload(ScheduleItem.module))
            .order_by(ScheduleItem.start_time)
        )
        if module_id:
            stmt = stmt.where(ScheduleItem.module_id == module_id)
        items = (await session.scalars(stmt)).all()
        lines = [f"📅 <b>{selected_day:%d.%m.%Y}</b>", ""]
        grouped: dict[tuple[Any, ...], list[ScheduleItem]] = defaultdict(list)
        for item in items:
            key = (
                item.module_id,
                item.start_time,
                item.end_time,
                item.title,
                item.location,
                item.format,
                item.comment,
            )
            grouped[key].append(item)

        for _, group_items in grouped.items():
            item = group_items[0]
            module_speakers = [
                await get_module_speaker(session, item.module_id, row.speaker_id)
                for row in group_items
                if row.speaker_id
            ]
            module_speakers = [ms for ms in module_speakers if ms]
            speaker_names = [row.speaker.full_name for row in group_items if row.speaker]
            setup_items = sorted({setup for ms in module_speakers for setup in (ms.setup or [])})
            total_amount = sum(ms.amount or 0 for ms in module_speakers)
            needs_flowers = any(ms.flower_required or ms.speaker.gender.lower().startswith("жен") for ms in module_speakers)
            lines += [
                f"<b>{time_label(item.start_time)}-{time_label(item.end_time)}</b>",
                (
                    "🎤 <b>Спикеры:</b> "
                    + ", ".join(f"<b>{html_escape(name)}</b>" for name in speaker_names)
                    if speaker_names
                    else "🎤 <b>Спикера нет</b>"
                ),
                f"Тема: {html_escape(item.title)}",
                f"Сетап: {html_escape(', '.join(setup_items) if setup_items else 'не указан')}",
            ]
            if total_amount:
                lines.append(f"Сумма: {money(total_amount)}")
            if needs_flowers:
                lines.append("🌸 Нужны цветы")
            lines.append("")
    await send_or_edit(callback, "\n".join(lines), kb_back(module_id or None))


async def recreate_reminders_for_module(session: AsyncSession, module_id: int, user_id: int) -> None:
    settings = await get_settings(session, user_id)
    tz = ZoneInfo(settings.timezone)
    for old in (await session.scalars(select(Reminder).where(Reminder.module_id == module_id, Reminder.is_sent == False))).all():
        await session.delete(old)
    items = (
        await session.scalars(
            select(ScheduleItem)
            .where(ScheduleItem.module_id == module_id, ScheduleItem.speaker_id.is_not(None))
            .options(selectinload(ScheduleItem.speaker))
        )
    ).all()
    for item in items:
        if not item.start_time:
            continue
        start_dt = datetime.combine(item.date, item.start_time).replace(tzinfo=tz).astimezone(ZoneInfo("UTC")).replace(tzinfo=None)
        end_dt = datetime.combine(item.date, item.end_time or item.start_time).replace(tzinfo=tz).astimezone(ZoneInfo("UTC")).replace(tzinfo=None)
        ms = await get_module_speaker(session, item.module_id, item.speaker_id)
        session.add(Reminder(module_id=item.module_id, speaker_id=item.speaker_id, schedule_item_id=item.id, reminder_type="performance", remind_at=start_dt - timedelta(minutes=settings.performance_reminder_minutes), target_user_id=user_id))
        if ms and ms.is_paid:
            session.add(Reminder(module_id=item.module_id, speaker_id=item.speaker_id, schedule_item_id=item.id, reminder_type="documents", remind_at=start_dt - timedelta(minutes=settings.docs_reminder_minutes), target_user_id=user_id))
        if ms and (ms.flower_required or ms.speaker.gender.lower().startswith("жен")) and item.end_time:
            session.add(Reminder(module_id=item.module_id, speaker_id=item.speaker_id, schedule_item_id=item.id, reminder_type="flowers", remind_at=end_dt - timedelta(minutes=settings.flowers_reminder_minutes), target_user_id=user_id))


async def send_due_reminders(bot: Bot) -> None:
    async with SessionLocal() as session:
        now = datetime.utcnow()
        reminders = (
            await session.scalars(
                select(Reminder)
                .where(Reminder.is_sent == False, Reminder.remind_at <= now)
                .order_by(Reminder.remind_at)
                .limit(25)
            )
        ).all()
        for rem in reminders:
            text = await build_reminder_text(session, rem)
            try:
                await bot.send_message(rem.target_user_id, text)
                rem.is_sent = True
                rem.sent_at = datetime.utcnow()
            except Exception as exc:
                log.warning("cannot send reminder %s: %s", rem.id, exc)
        await session.commit()


async def build_reminder_text(session: AsyncSession, rem: Reminder) -> str:
    module = await session.get(Module, rem.module_id)
    speaker = await session.get(Speaker, rem.speaker_id) if rem.speaker_id else None
    item = await session.get(ScheduleItem, rem.schedule_item_id) if rem.schedule_item_id else None
    ms = await get_module_speaker(session, rem.module_id, rem.speaker_id) if rem.speaker_id else None
    if rem.reminder_type == "documents" and ms:
        ds = ms.document_status
        lines = [
            "📄 <b>Проверь документы по платному спикеру</b>",
            "",
            f"👤 {html_escape(speaker.full_name)}",
            f"⏰ Выступление через 30 минут",
            f"💵 Сумма: {money(ms.amount)}",
            f"📄 Тип оформления: {html_escape(ms.contract_type)}",
            "",
            "Статус документов:",
        ]
        for field, title in DOC_FIELDS:
            lines.append(f"{'✅' if ds and getattr(ds, field) else '❌'} {title}")
        lines.append(f"\nПрогресс: {doc_progress(ds)}%")
        return "\n".join(lines)
    if rem.reminder_type == "flowers":
        return "\n".join([
            "🌸 <b>Через 30 минут нужно вынести цветы</b>",
            "",
            f"👤 Спикер: {html_escape(speaker.full_name if speaker else '')}",
            f"📁 Модуль: {html_escape(module.name if module else '')}",
            f"⏰ Выступление заканчивается в {time_label(item.end_time) if item else '--:--'}",
            f"📍 Зал: {html_escape(item.location if item else '')}",
            "",
            "Не забудь подготовить букет и передать его в зал.",
        ])
    return "\n".join([
        "🔔 <b>Через 40 минут выступает спикер</b>",
        "",
        f"👤 {html_escape(speaker.full_name if speaker else '')}",
        f"📁 Модуль: {html_escape(module.name if module else '')}",
        f"⏰ Время: {time_label(item.start_time) if item else '--:--'}-{time_label(item.end_time) if item else '--:--'}",
        f"🎤 Тема: {html_escape(item.title if item else (ms.topic if ms else ''))}",
        f"📍 Зал: {html_escape(item.location if item else '')}",
        "",
        "⚙️ Сетап:",
        *(f"- {html_escape(x)}" for x in (ms.setup if ms and ms.setup else ["не указан"])),
        "",
        f"📞 Контакт: {html_escape(speaker.phone if speaker else '')}",
        f"💬 Telegram: {html_escape(speaker.telegram if speaker else '')}",
        f"💰 Сумма: {money(ms.amount if ms else 0)}",
    ])


async def post_module_document_reminders(bot: Bot) -> None:
    async with SessionLocal() as session:
        users = (await session.scalars(select(AppSetting))).all()
        today = datetime.utcnow().date()
        for settings in users:
            if not settings.post_module_reminders_enabled:
                continue
            modules = (
                await session.scalars(
                    select(Module).where(Module.owner_user_id == settings.user_id, Module.end_date.is_not(None), Module.end_date < today)
                )
            ).all()
            for module in modules:
                if (today - module.end_date).days % max(settings.post_module_interval_days, 1) != 0:
                    continue
                links = (
                    await session.scalars(
                        select(ModuleSpeaker)
                        .where(ModuleSpeaker.module_id == module.id, ModuleSpeaker.is_paid == True)
                        .options(selectinload(ModuleSpeaker.speaker), selectinload(ModuleSpeaker.document_status))
                    )
                ).all()
                for ms in links:
                    if doc_progress(ms.document_status) >= 100:
                        continue
                    missing = [title for field, title in DOC_FIELDS if not ms.document_status or not getattr(ms.document_status, field)]
                    text = "\n".join([
                        "📄 <b>Не закрыты документы по спикеру</b>",
                        "",
                        f"👤 {html_escape(ms.speaker.full_name)}",
                        f"📁 Модуль: {html_escape(module.name)}",
                        f"💵 Сумма: {money(ms.amount)}",
                        f"📄 Тип оформления: {html_escape(ms.contract_type)}",
                        "",
                        "Не закрыто:",
                        *(f"- {html_escape(x)}" for x in missing),
                        "",
                        f"Прогресс: {doc_progress(ms.document_status)}%",
                    ])
                    try:
                        await bot.send_message(settings.user_id, text)
                    except Exception as exc:
                        log.warning("post module reminder failed: %s", exc)
        await session.commit()


@router.callback_query(F.data == "reminders")
async def reminders_view(callback: CallbackQuery) -> None:
    async with SessionLocal() as session:
        reminders = (
            await session.scalars(
                select(Reminder)
                .where(Reminder.target_user_id == callback.from_user.id, Reminder.is_sent == False)
                .order_by(Reminder.remind_at)
                .limit(20)
            )
        ).all()
    lines = ["🔔 <b>Ближайшие напоминания</b>", ""]
    for rem in reminders:
        lines.append(f"{rem.remind_at:%d.%m.%Y %H:%M} UTC | {rem.reminder_type}")
    await send_or_edit(callback, "\n".join(lines) if reminders else "Активных напоминаний пока нет.", kb_main())


@router.callback_query(F.data.startswith("recreate_reminders:"))
async def recreate_reminders_callback(callback: CallbackQuery) -> None:
    ms_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        ms = await session.get(ModuleSpeaker, ms_id)
        if ms:
            await recreate_reminders_for_module(session, ms.module_id, callback.from_user.id)
            await session.commit()
    await callback.answer("Напоминания обновлены", show_alert=True)


@router.callback_query(F.data == "settings")
async def settings_view(callback: CallbackQuery) -> None:
    async with SessionLocal() as session:
        s = await get_settings(session, callback.from_user.id)
        await session.commit()
    text = "\n".join([
        "⚙️ <b>Настройки</b>",
        "",
        f"Часовой пояс: {html_escape(s.timezone)}",
        f"Напоминание перед выступлением: {s.performance_reminder_minutes} мин.",
        f"Напоминание о документах: {s.docs_reminder_minutes} мин.",
        f"Напоминание о цветах: {s.flowers_reminder_minutes} мин. до конца",
        f"ИИ-помощник: {'включен' if s.ai_enabled else 'выключен'}",
        f"Напоминания после модуля: {'включены' if s.post_module_reminders_enabled else 'выключены'}",
        f"Интервал после модуля: каждые {s.post_module_interval_days} дн.",
    ])
    kb = InlineKeyboardBuilder()
    buttons = [
        ("Часовой пояс", "settings_edit:timezone"),
        ("Минуты до выступления", "settings_edit:performance_minutes"),
        ("Документы: 20/30", "settings_docs_toggle"),
        ("Минуты до цветов", "settings_edit:flowers_minutes"),
        ("ИИ вкл/выкл", "settings_toggle_ai"),
        ("После модуля вкл/выкл", "settings_toggle_post"),
        ("Интервал после модуля", "settings_edit:post_interval"),
        ("👥 Доступ", "access_menu"),
        ("🏠 Главное меню", "menu"),
    ]
    for t, d in buttons:
        kb.button(text=t, callback_data=d)
    kb.adjust(2, 2, 2, 1, 1, 1)
    await send_or_edit(callback, text, kb)


@router.callback_query(F.data == "access_menu")
async def access_menu(callback: CallbackQuery) -> None:
    async with SessionLocal() as session:
        grants = (
            await session.scalars(
                select(AccessGrant).where(AccessGrant.owner_user_id == callback.from_user.id).order_by(AccessGrant.created_at)
            )
        ).all()
    lines = ["👥 <b>Доступ к вашим модулям</b>", ""]
    if grants:
        for i, grant in enumerate(grants, 1):
            resolved = f"ID {grant.allowed_user_id}" if grant.allowed_user_id else "ждет /start"
            allowed = ", ".join(title for key, title in ACCESS_PERMISSIONS if key in normalize_access_permissions(grant.permissions))
            lines.append(f"{i}. @{html_escape(grant.username)} - {resolved}")
            lines.append(f"   Кнопки: {html_escape(allowed)}")
    else:
        lines.append("Пока никому не выдан доступ.")
    lines += ["", "Добавьте username без ссылки. Человек должен написать боту /start."]
    kb = InlineKeyboardBuilder()
    kb.button(text="➕ Добавить username", callback_data="access_add")
    for grant in grants:
        kb.button(text=f"⚙️ Права @{grant.username}", callback_data=f"access_permissions:{grant.id}")
        kb.button(text=f"🗑 @{grant.username}", callback_data=f"access_delete:{grant.id}")
    kb.button(text="⬅️ Настройки", callback_data="settings")
    kb.adjust(1)
    await send_or_edit(callback, "\n".join(lines), kb)


@router.callback_query(F.data == "access_add")
async def access_add(callback: CallbackQuery, state: FSMContext) -> None:
    await state.set_state(AccessEdit.username)
    await callback.message.edit_text("Введите Telegram username человека, например: @username")
    await callback.answer()


@router.message(AccessEdit.username)
async def access_add_save(message: Message, state: FSMContext) -> None:
    username = normalize_username(message.text)
    if not username:
        await message.answer("Не понял username. Отправьте в формате @username.")
        return
    async with SessionLocal() as session:
        existing = await session.scalar(
            select(AccessGrant).where(AccessGrant.owner_user_id == message.from_user.id, AccessGrant.username == username)
        )
        if not existing:
            grant = AccessGrant(owner_user_id=message.from_user.id, username=username, permissions=DEFAULT_GRANT_PERMISSIONS.copy())
            session.add(grant)
            await log_action(session, message.from_user.id, "grant_access", "access_grant", None, new={"username": username})
        await session.commit()
    await state.clear()
    await message.answer(f"✅ Доступ выдан @{username}. Пусть человек напишет боту /start.", reply_markup=kb_main().as_markup())


@router.callback_query(F.data.startswith("access_delete:"))
async def access_delete(callback: CallbackQuery) -> None:
    grant_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        grant = await session.get(AccessGrant, grant_id)
        if grant and grant.owner_user_id == callback.from_user.id:
            await log_action(session, callback.from_user.id, "revoke_access", "access_grant", grant.id, old={"username": grant.username})
            await session.delete(grant)
            await session.commit()
    await access_menu(callback)


@router.callback_query(F.data.startswith("access_permissions:"))
async def access_permissions_menu(callback: CallbackQuery) -> None:
    grant_id = int(callback.data.split(":")[1])
    async with SessionLocal() as session:
        grant = await session.get(AccessGrant, grant_id)
        if not grant or grant.owner_user_id != callback.from_user.id:
            await callback.answer("Доступ не найден.", show_alert=True)
            return
        selected = set(normalize_access_permissions(grant.permissions))
        lines = [
            f"⚙️ <b>Кнопки для @{html_escape(grant.username)}</b>",
            "",
            "Нажимайте на разделы, чтобы включать или выключать их.",
            "",
        ]
        kb = InlineKeyboardBuilder()
        for key, title in ACCESS_PERMISSIONS:
            enabled = key in selected
            lines.append(f"{'✅' if enabled else '❌'} {html_escape(title)}")
            kb.button(text=f"{'✅' if enabled else '❌'} {title}", callback_data=f"access_toggle:{grant.id}:{key}")
        kb.button(text="⬅️ Доступ", callback_data="access_menu")
        kb.adjust(1)
    await send_or_edit(callback, "\n".join(lines), kb)


@router.callback_query(F.data.startswith("access_toggle:"))
async def access_toggle(callback: CallbackQuery) -> None:
    _, grant_id_raw, key = callback.data.split(":")
    known = {permission_key for permission_key, _title in ACCESS_PERMISSIONS}
    if key not in known:
        await callback.answer("Неизвестная кнопка.", show_alert=True)
        return
    grant_id = int(grant_id_raw)
    async with SessionLocal() as session:
        grant = await session.get(AccessGrant, grant_id)
        if grant and grant.owner_user_id == callback.from_user.id:
            permissions = set(normalize_access_permissions(grant.permissions))
            old = {"permissions": sorted(permissions)}
            if key in permissions:
                permissions.remove(key)
            else:
                permissions.add(key)
            grant.permissions = sorted(permissions)
            await log_action(session, callback.from_user.id, "toggle_access_permission", "access_grant", grant.id, old=old, new={"permissions": grant.permissions})
            await session.commit()
    await access_permissions_menu(callback)


@router.callback_query(F.data == "settings_docs_toggle")
async def settings_docs_toggle(callback: CallbackQuery) -> None:
    async with SessionLocal() as session:
        s = await get_settings(session, callback.from_user.id)
        s.docs_reminder_minutes = 20 if s.docs_reminder_minutes == 30 else 30
        await session.commit()
    await settings_view(callback)


@router.callback_query(F.data.in_({"settings_toggle_ai", "settings_toggle_post"}))
async def settings_toggle(callback: CallbackQuery) -> None:
    async with SessionLocal() as session:
        s = await get_settings(session, callback.from_user.id)
        if callback.data == "settings_toggle_ai":
            s.ai_enabled = not s.ai_enabled
        else:
            s.post_module_reminders_enabled = not s.post_module_reminders_enabled
        await session.commit()
    await settings_view(callback)


@router.callback_query(F.data.startswith("settings_edit:"))
async def settings_edit(callback: CallbackQuery, state: FSMContext) -> None:
    field = callback.data.split(":")[1]
    await state.update_data(field=field)
    await state.set_state(getattr(SettingsEdit, field))
    await callback.message.edit_text("Введите новое значение:")
    await callback.answer()


@router.message(SettingsEdit.timezone)
@router.message(SettingsEdit.performance_minutes)
@router.message(SettingsEdit.flowers_minutes)
@router.message(SettingsEdit.post_interval)
async def settings_save(message: Message, state: FSMContext) -> None:
    data = await state.get_data()
    field = data["field"]
    value = message.text.strip()
    async with SessionLocal() as session:
        s = await get_settings(session, message.from_user.id)
        if field == "timezone":
            ZoneInfo(value)
            s.timezone = value
        else:
            setattr(s, field, max(1, int(re.sub(r"\D", "", value) or 1)))
        await session.commit()
    await state.clear()
    await message.answer("✅ Настройки сохранены.", reply_markup=kb_main().as_markup())


async def parse_text_command(message_text: str) -> dict[str, Any] | None:
    if not openai_client or not OPENAI_API_KEY:
        return heuristic_text_command(message_text)
    schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "intent": {"type": "string", "enum": ["create_speaker", "update_speaker_payment", "update_document_status", "update_setup", "update_schedule_time", "unknown"]},
            "speaker_name": {"type": "string"},
            "module_name": {"type": "string"},
            "fields": {"type": "object", "additionalProperties": True},
            "needs_confirmation": {"type": "boolean"},
            "summary": {"type": "string"},
        },
        "required": ["intent", "speaker_name", "module_name", "fields", "needs_confirmation", "summary"],
    }
    try:
        response = await openai_client.responses.create(
            model=OPENAI_MODEL,
            input=[
                {"role": "system", "content": "Преобразуй сообщение менеджера мероприятий в структурированную команду для бота. Важные изменения помечай needs_confirmation=true."},
                {"role": "user", "content": message_text},
            ],
            text={"format": {"type": "json_schema", "name": "bot_command", "schema": schema, "strict": True}},
        )
        return json.loads(response.output_text)
    except Exception as exc:
        log.exception("AI command parse failed: %s", exc)
        return heuristic_text_command(message_text)


def heuristic_text_command(text: str) -> dict[str, Any] | None:
    low = text.lower()
    amount = re.search(r"(\d[\d\s]{3,})", text)
    name = ""
    m = re.search(r"(иванов[а-яё\s]*|петров[а-яё\s]*|смирнов[а-яё\s]*|[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+){0,2})", text)
    if m:
        name = m.group(1).strip()
    if "сумм" in low or amount:
        return {"intent": "update_speaker_payment", "speaker_name": name, "module_name": "", "fields": {"amount": int(re.sub(r"\D", "", amount.group(1))) if amount else 0}, "needs_confirmation": True, "summary": "изменить сумму"}
    if "подписан" in low or "счет" in low or "чек" in low or "договор" in low:
        fields = {}
        if "подпис" in low:
            fields["contract_act_signed"] = True
        if "счет" in low or "чек" in low:
            fields["invoice_or_receipt_received"] = True
        if "соглас" in low:
            fields["contract_act_approved"] = True
        return {"intent": "update_document_status", "speaker_name": name, "module_name": "", "fields": fields, "needs_confirmation": True, "summary": "обновить документы"}
    if "цвет" in low or "кликер" in low or "презентац" in low or "микрофон" in low:
        return {"intent": "update_setup", "speaker_name": name, "module_name": "", "fields": {"setup_add": split_setup(text)}, "needs_confirmation": True, "summary": "обновить сетап"}
    if "добав" in low and name:
        return {"intent": "create_speaker", "speaker_name": name, "module_name": "", "fields": {}, "needs_confirmation": True, "summary": "добавить спикера"}
    return None


async def apply_text_command(session: AsyncSession, user_id: int, command: dict[str, Any]) -> str:
    intent = command.get("intent")
    speaker = await find_speaker(session, user_id, command.get("speaker_name", ""))
    module = await find_module(session, user_id, command.get("module_name", ""))
    fields = command.get("fields") or {}
    if intent == "create_speaker":
        if not module:
            return "Не нашел модуль. Создайте модуль или укажите его название."
        speaker = Speaker(full_name=command.get("speaker_name") or fields.get("full_name") or "Новый спикер", gender=fields.get("gender", ""), phone=fields.get("phone", ""), telegram=fields.get("telegram", ""), email=fields.get("email", ""), owner_user_id=user_id)
        session.add(speaker)
        await session.flush()
        ms = await ensure_module_speaker(session, module, speaker, topic=fields.get("topic", ""), setup=split_setup(fields.get("setup", [])), is_paid=bool(fields.get("amount")), amount=float(fields.get("amount") or 0), contract_type=fields.get("contract_type", ""))
        await log_action(session, user_id, "ai_create_speaker", "speaker", speaker.id, new=fields)
        await session.commit()
        return f"✅ Добавил спикера: {speaker.full_name}. Карточка доступна в модуле «{module.name}»."
    if not speaker:
        return "Не нашел спикера. Уточните ФИО."
    if not module:
        link = await session.scalar(
            select(ModuleSpeaker)
            .where(ModuleSpeaker.speaker_id == speaker.id)
            .options(selectinload(ModuleSpeaker.module))
        )
        module = link.module if link else None
    if not module:
        return "Не нашел модуль спикера."
    ms = await get_module_speaker(session, module.id, speaker.id)
    if intent == "update_speaker_payment":
        old = {"amount": ms.amount}
        ms.amount = float(fields.get("amount") or 0)
        ms.is_paid = ms.amount > 0
        if fields.get("contract_type"):
            ms.contract_type = fields["contract_type"]
        if ms.is_paid and not ms.document_status:
            session.add(DocumentStatus(module_speaker_id=ms.id))
        await log_action(session, user_id, "ai_update_payment", "module_speaker", ms.id, old=old, new=fields)
        await session.commit()
        return f"✅ Обновил сумму для {speaker.full_name}: {money(ms.amount)}."
    if intent == "update_document_status":
        if not ms.document_status:
            ms.document_status = DocumentStatus(module_speaker_id=ms.id)
        for field in fields:
            if field in dict(DOC_FIELDS):
                setattr(ms.document_status, field, bool(fields[field]))
        doc_progress(ms.document_status)
        await log_action(session, user_id, "ai_update_docs", "document_status", ms.document_status.id, new=fields)
        await session.commit()
        return f"✅ Обновил документы для {speaker.full_name}. Прогресс: {ms.document_status.progress_percent}%."
    if intent == "update_setup":
        additions = split_setup(fields.get("setup_add") or fields.get("setup") or [])
        ms.setup = sorted(set((ms.setup or []) + additions))
        ms.flower_required = ms.flower_required or "цветы" in ms.setup or bool(fields.get("flower_required"))
        await log_action(session, user_id, "ai_update_setup", "module_speaker", ms.id, new=fields)
        await session.commit()
        return f"✅ Обновил сетап для {speaker.full_name}: {', '.join(ms.setup)}."
    if intent == "update_schedule_time":
        item = await session.scalar(select(ScheduleItem).where(ScheduleItem.speaker_id == speaker.id, ScheduleItem.module_id == module.id))
        if not item:
            return "Не нашел слот расписания для этого спикера."
        item.start_time = parse_time(str(fields.get("start_time", ""))) or item.start_time
        item.end_time = parse_time(str(fields.get("end_time", ""))) or item.end_time
        await recreate_reminders_for_module(session, module.id, user_id)
        await log_action(session, user_id, "ai_update_schedule", "schedule_item", item.id, new=fields)
        await session.commit()
        return f"✅ Перенес выступление {speaker.full_name} на {time_label(item.start_time)}."
    return "Не понял действие. Попробуйте написать чуть конкретнее."


@router.callback_query(F.data.in_({"ai_confirm", "ai_cancel"}))
async def ai_confirm(callback: CallbackQuery) -> None:
    command = PENDING_AI_ACTIONS.pop(callback.from_user.id, None)
    permissions = await effective_permissions(callback.from_user.id, callback.from_user.username)
    if callback.data == "ai_cancel" or not command:
        await callback.message.edit_text("Отменено.", reply_markup=kb_main(permissions).as_markup())
        await callback.answer()
        return
    async with SessionLocal() as session:
        result = await apply_text_command(session, callback.from_user.id, command)
    await callback.message.edit_text(result, reply_markup=kb_main(permissions).as_markup())
    await callback.answer()


@router.message(F.text)
async def natural_text(message: Message, state: FSMContext) -> None:
    current = await state.get_state()
    if current:
        return
    async with SessionLocal() as session:
        settings = await get_settings(session, message.from_user.id)
        permissions = await access_permissions(session, message.from_user.id, message.from_user.username)
        if not settings.ai_enabled:
            await message.answer("ИИ-помощник выключен в настройках.", reply_markup=kb_main(permissions).as_markup())
            return
    command = await parse_text_command(message.text)
    if not command or command.get("intent") == "unknown":
        permissions = await effective_permissions(message.from_user.id, message.from_user.username)
        await message.answer("Не понял задачу. Можно нажать кнопку в меню или написать подробнее.", reply_markup=kb_main(permissions).as_markup())
        return
    if command.get("needs_confirmation", True):
        PENDING_AI_ACTIONS[message.from_user.id] = command
        kb = InlineKeyboardBuilder()
        kb.button(text="✅ Да, выполнить", callback_data="ai_confirm")
        kb.button(text="❌ Отмена", callback_data="ai_cancel")
        kb.adjust(1)
        await message.answer(f"Я понял задачу:\n\n{html_escape(command.get('summary') or json.dumps(command, ensure_ascii=False))}\n\nПодтвердить изменение?", reply_markup=kb.as_markup())
    else:
        async with SessionLocal() as session:
            result = await apply_text_command(session, message.from_user.id, command)
        permissions = await effective_permissions(message.from_user.id, message.from_user.username)
        await message.answer(result, reply_markup=kb_main(permissions).as_markup())


async def init_db() -> None:
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
        dialect = conn.dialect.name
        default_permissions = json.dumps(DEFAULT_GRANT_PERMISSIONS)
        if dialect == "sqlite":
            rows = await conn.execute(sql_text("PRAGMA table_info(access_grants)"))
            columns = {row[1] for row in rows.fetchall()}
            if "permissions" not in columns:
                await conn.execute(sql_text("ALTER TABLE access_grants ADD COLUMN permissions JSON DEFAULT '[]'"))
                await conn.execute(sql_text(f"UPDATE access_grants SET permissions = '{default_permissions}'"))
        elif dialect in {"postgresql", "postgres"}:
            await conn.execute(sql_text("ALTER TABLE access_grants ADD COLUMN IF NOT EXISTS permissions JSON DEFAULT '[]'"))


async def main() -> None:
    if not BOT_TOKEN:
        raise RuntimeError("BOT_TOKEN не задан. Создайте .env и укажите BOT_TOKEN.")
    await init_db()
    bot = Bot(BOT_TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))
    dp = Dispatcher(storage=MemoryStorage())
    dp.message.middleware(AccessMiddleware())
    dp.callback_query.middleware(AccessMiddleware())
    dp.include_router(router)
    scheduler = AsyncIOScheduler(timezone="UTC")
    scheduler.add_job(send_due_reminders, "interval", seconds=60, args=[bot], id="send_due_reminders")
    scheduler.add_job(post_module_document_reminders, "interval", hours=6, args=[bot], id="post_module_document_reminders")
    scheduler.start()
    log.info("Bot started")
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())
